#!/usr/bin/env python3
"""Generate Pipelines_and_BIDS_Documentation.docx (requires python-docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)


def add_link_paragraph(doc: Document, label: str, url: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ")
    p.add_run(url)


def add_field_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val


def add_reference_appendix(doc: Document) -> None:
    """Full list of resources (mirrors pipelines/pipelines.yaml reference_resources)."""
    doc.add_heading("8. References and resources (complete list)", level=1)
    doc.add_paragraph(
        "Consolidated list of normative and background links used for recent pipelines "
        "(BIDS layout and HippUnfold). Keep in sync with pipelines/pipelines.yaml under key "
        "reference_resources."
    )

    doc.add_heading("8.1 BIDS specification", level=2)
    bids_refs = [
        (
            "Dataset description (dataset_description.json, README, derivatives)",
            "https://bids-specification.readthedocs.io/en/stable/modality-agnostic-files/dataset-description.html",
        ),
        (
            "Data summary files (participants, samples, scans, sessions)",
            "https://bids-specification.readthedocs.io/en/stable/modality-agnostic-files/data-summary-files.html",
        ),
        (
            "Scans file (_scans.tsv / _scans.json)",
            "https://bids-specification.readthedocs.io/en/stable/modality-agnostic-files/data-summary-files.html#scans-file",
        ),
        (
            "Common principles (tabular files, BIDS URIs, units)",
            "https://bids-specification.readthedocs.io/en/stable/common-principles.html",
        ),
        (
            "Glossary",
            "https://bids-specification.readthedocs.io/en/stable/glossary.html",
        ),
        (
            "BIDS starter kit (README templates and examples)",
            "https://github.com/bids-standard/bids-starter-kit",
        ),
    ]
    for title, url in bids_refs:
        add_link_paragraph(doc, title, url)

    doc.add_heading("8.2 HippUnfold", level=2)
    hf_refs = [
        ("Project landing (overview, workflow, publications)", "https://hippunfold.khanlab.ca/en/latest/"),
        ("Full documentation (Read the Docs)", "https://hippunfold.readthedocs.io/en/latest/"),
        ("Source repository", "https://github.com/khanlab/hippunfold"),
        ("Docker Hub (khanlab/hippunfold)", "https://hub.docker.com/r/khanlab/hippunfold"),
        ("hippunfold_toolbox", "https://github.com/jordandekraker/hippunfold_toolbox"),
        ("Hippo_Spin_Testing", "https://github.com/Bradley-Karat/Hippo_Spin_Testing"),
        ("SynthSeg (optional contrast-agnostic segmentation)", "https://github.com/BBillot/SynthSeg"),
    ]
    for title, url in hf_refs:
        add_link_paragraph(doc, title, url)

    doc.add_heading("8.3 Publications (HippUnfold)", level=2)
    pub_refs = [
        ("DeKraker et al., 2022 — HippUnfold methods (cite for any version)", "https://doi.org/10.7554/eLife.77945"),
        (
            "DeKraker et al., 2023 — Unfolded registration / multihist7 (cite for HippUnfold >= 1.3.0)",
            "https://doi.org/10.7554/eLife.88404.3",
        ),
    ]
    for title, url in pub_refs:
        add_link_paragraph(doc, title, url)


def build() -> Document:
    doc = Document()

    add_title(doc, "Neuroimaging Pipelines and BIDS Documentation")

    doc.add_paragraph(
        "This document registers analysis pipelines used in the project, summarizes "
        "Brain Imaging Data Structure (BIDS) requirements relevant to raw and derivative "
        "datasets, and documents modality-specific metadata practices—including the "
        "session-level scans file—so that datasets remain interoperable and citable."
    )

    doc.add_heading("1. Normative references", level=1)
    add_link_paragraph(
        doc,
        "BIDS dataset description (dataset_description.json, README, derivatives)",
        "https://bids-specification.readthedocs.io/en/stable/modality-agnostic-files/dataset-description.html",
    )
    add_link_paragraph(
        doc,
        "BIDS data summary files (participants, samples, scans, sessions)",
        "https://bids-specification.readthedocs.io/en/stable/modality-agnostic-files/data-summary-files.html",
    )
    add_link_paragraph(
        doc,
        "BIDS scans file (subsection on _scans.tsv)",
        "https://bids-specification.readthedocs.io/en/stable/modality-agnostic-files/data-summary-files.html#scans-file",
    )
    add_link_paragraph(
        doc,
        "HippUnfold (project overview and links to full documentation)",
        "https://hippunfold.khanlab.ca/en/latest/",
    )
    add_link_paragraph(
        doc,
        "HippUnfold Read the Docs",
        "https://hippunfold.readthedocs.io/en/latest/",
    )

    doc.add_heading("2. Pipeline registry", level=1)
    doc.add_paragraph(
        "The machine-readable registry is stored as pipelines/pipelines.yaml alongside this documentation. "
        "It currently includes HippUnfold for hippocampal unfolding and derivative outputs."
    )

    doc.add_heading("2.1 HippUnfold", level=2)
    doc.add_paragraph(
        "HippUnfold models the topological folding structure of the human hippocampus and "
        "computationally unfolds it. Typical use cases include visualization, topologically constrained "
        "intersubject registration, parcellation against an unfolded atlas, morphometry (thickness, "
        "surface area, curvature, gyrification), and quantitative mapping (for example qT1 mapped to "
        "midthickness surfaces and laminar profiles)."
    )
    doc.add_paragraph(
        "Version 1.3.x introduces unfolded-space registration to a reference atlas harmonized across "
        "seven ground-truth histology samples; the legacy workflow can be selected with --atlas bigbrain "
        "or --no-unfolded-reg. Containers from 1.3 onward download nnU-Net models at run time rather than "
        "shipping them in the image."
    )
    add_link_paragraph(doc, "Source code", "https://github.com/khanlab/hippunfold")
    doc.add_paragraph("Related tools:", style="List Bullet")
    doc.add_paragraph(
        "hippunfold_toolbox — plotting, mapping fMRI/DWI, surface tools: "
        "https://github.com/jordandekraker/hippunfold_toolbox",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Hippo_Spin_Testing — spin tests in unfolded space: "
        "https://github.com/Bradley-Karat/Hippo_Spin_Testing",
        style="List Bullet",
    )
    doc.add_paragraph("Citations:", style="List Bullet")
    doc.add_paragraph(
        "DeKraker et al., eLife 2022 (HippUnfold methods): https://doi.org/10.7554/eLife.77945",
        style="List Bullet",
    )
    doc.add_paragraph(
        "DeKraker et al., eLife 2023 (unfolded registration; cite if using HippUnfold ≥ 1.3.0): "
        "https://doi.org/10.7554/eLife.88404.3",
        style="List Bullet",
    )

    doc.add_heading("3. BIDS: dataset-level description (all modalities)", level=1)
    doc.add_paragraph(
        "Every BIDS dataset must include dataset_description.json at the dataset root. "
        "Required keys are Name and BIDSVersion. Recommended fields include DatasetType (raw, derivative, "
        "or study), License, Authors, Keywords, GeneratedBy (recommended for provenance), and "
        "SourceDatasets when derivatives are derived from identifiable sources."
    )
    doc.add_paragraph(
        "Derivative datasets stored under derivatives/<pipeline>/ must also include dataset_description.json "
        "at that derivative root. For derivatives, GeneratedBy is required: each object should include "
        "Name (pipeline name), and should include Version, Description, CodeURL, and Container when applicable. "
        "If the derivative folder is nested under the raw dataset, the first GeneratedBy Name should be a "
        "substring of the derivative directory name, per BIDS."
    )

    add_field_table(
        doc,
        ["Key", "Raw requirement", "Derivative requirement"],
        [
            ["Name", "Required", "Required"],
            ["BIDSVersion", "Required", "Required"],
            ["DatasetType", "Recommended (default raw)", "Should be derivative"],
            ["GeneratedBy", "Recommended", "Required"],
            ["SourceDatasets", "Recommended if applicable", "Recommended"],
        ],
    )
    doc.add_paragraph()

    doc.add_heading("4. BIDS: data summary files and the scans file", level=1)
    doc.add_paragraph(
        "Data summary files describe participants, samples (when sample- entities exist), sessions, "
        "and—critically for per-session acquisition metadata—the scans file."
    )

    doc.add_heading("4.1 participants.tsv (and participants.json)", level=2)
    doc.add_paragraph(
        "If present, participants.tsv must include participant_id as the first column, with one row per "
        "participant. Common optional columns include age, sex, handedness, species, strain. "
        "Sidecar participants.json should define columns and levels for interoperability."
    )

    doc.add_heading("4.2 Sessions file (multi-session studies)", level=2)
    doc.add_paragraph(
        "When multiple sessions exist, sub-<label>_sessions.tsv may describe session-level variables. "
        "It must include session_id as the first column, one row per session."
    )

    doc.add_heading("4.3 Scans file (_scans.tsv and _scans.json)", level=2)
    doc.add_paragraph(
        "Template: sub-<label>/[ses-<label>/]sub-<label>[_ses-<label>]_scans.tsv and optional "
        "sub-<label>[_ses-<label>]_scans.json. Purpose: describe timing and other properties of each "
        "recording file within one session. Each listed file should generally have exactly one row."
    )
    doc.add_paragraph(
        "For multi-part file formats, list only the file passed to analysis software (for example the "
        ".vhdr for BrainVision). For entity-linked collections (for example multiple echoes or splits), "
        "use one row per file as required by BIDS."
    )
    add_field_table(
        doc,
        ["Column", "Requirement", "Description"],
        [
            ["filename", "Required", "Relative paths; first column; unique per row"],
            ["acq_time", "Optional", "ISO8601 datetime of first data point; identical for linked splits"],
            ["HED", "Optional", "Hierarchical Event Descriptor tags if used"],
            ["Additional columns", "Optional", "Document in _scans.json"],
        ],
    )
    doc.add_paragraph()

    doc.add_heading("5. Modality-specific BIDS documentation", level=1)
    doc.add_paragraph(
        "Below: what to capture in dataset_description / sidecars / scans for each major imaging modality. "
        "Always place imaging data under the correct BIDS modality folder (anat, func, dwi, fmap, pet, etc.) "
        "and use the appropriate suffix and entities (task, acq, ce, rec, run, echo, part, etc.)."
    )

    modalities = [
        (
            "5.1 Structural MRI (anat)",
            [
                "Primary HippUnfold input: typically T1w (suffix T1w).",
                "Sidecar JSON must include required DICOM fields where applicable: RepetitionTime, EchoTime, "
                "FlipAngle, ImagingFrequency, Manufacturer, ManufacturersModelName, and slice timing fields as required by BIDS.",
                "In _scans.tsv, list anat/<filename>.nii.gz with acq_time if you track session timing.",
                "For quantitative maps (T1map, T2map, etc.), follow qMRI BIDS conventions and entity-linked collections.",
            ],
        ),
        (
            "5.2 Functional MRI (func)",
            [
                "Files: sub-*_task-*_bold.nii.gz with sub-*_task-*_bold.json.",
                "Record TaskName, RepetitionTime, SliceTiming, PhaseEncodingDirection, EchoTime for each run.",
                "_scans.tsv: one row per bold file (or per part of split collections per BIDS rules).",
                "After HippUnfold, functional data may be sampled to hippocampal surfaces using hippunfold_toolbox; "
                "document such steps in derivative dataset_description GeneratedBy or in code/ README.",
            ],
        ),
        (
            "5.3 Diffusion MRI (dwi)",
            [
                "Include dwi NIfTI, bval, bvec, and JSON with PhaseEncodingDirection, TotalReadoutTime, etc.",
                "_scans.tsv lists each dwi file (or collection) per session.",
                "Use for connectivity analyses complementary to HippUnfold surfaces when applicable.",
            ],
        ),
        (
            "5.4 Field maps (fmap)",
            [
                "GRE field maps, SE-EPI PEPOLAR, or TB1EPI etc. per BIDS fmap rules.",
                "Link to func/dwi via IntendedFor in JSON sidecars.",
            ],
        ),
        (
            "5.5 PET and other modalities",
            [
                "Follow BIDS PET or modality-specific extension rules; include radiotracer and timing metadata in sidecars.",
                "List each primary image in _scans.tsv for the session.",
            ],
        ),
    ]
    for title, bullets in modalities:
        doc.add_heading(title, level=2)
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("6. Derivative layout for HippUnfold", level=1)
    doc.add_paragraph(
        "Store HippUnfold outputs under derivatives/hippunfold/ (or another name containing hippunfold so "
        "that GeneratedBy Name matches BIDS derivative rules). Include dataset_description.json with "
        "DatasetType derivative and GeneratedBy describing HippUnfold version and container. "
        "Reference the raw BIDS dataset in SourceDatasets (URI, DOI, or version string)."
    )

    doc.add_heading("7. Example files shipped with this documentation", level=1)
    doc.add_paragraph(
        "Example JSON and TSV templates are provided under pipelines/bids/: "
        "dataset_description.raw.example.json, dataset_description.hippunfold_derivative.example.json, "
        "sub-01_ses-01_scans.example.tsv, and sub-01_ses-01_scans.example.json."
    )

    add_reference_appendix(doc)

    doc.add_heading("9. Document control", level=1)
    doc.add_paragraph(
        "Generated programmatically for consistency with BIDS 1.x and HippUnfold public documentation. "
        "Update pipelines/pipelines.yaml (including reference_resources) and regenerate this file when "
        "pipelines or BIDS versions change."
    )

    return doc


def main() -> None:
    out_dir = Path(__file__).resolve().parent
    out_path = out_dir / "Pipelines_and_BIDS_Documentation.docx"
    build().save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
