#!/usr/bin/env python3
"""Generate thought_experiment.docx from research-planning conversation notes."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text):
    doc.add_paragraph(text)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = Document()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Thought experiment")
    r.bold = True
    r.font.size = Pt(18)
    doc.add_paragraph()
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.add_run(
        "Documentation synthesizing discussion on AI/automation-inspired imaging standards, "
        "CIDUR_data, NeuroImage targeting, evidence, and references."
    ).italic = True
    doc.add_paragraph()

    h(doc, "1. Context and project framing", level=1)
    p(
        doc,
        "The underlying research direction is to propose AI- and automation-inspired acquisition "
        "and export standards, with Data Birth Integrity (DBI) as a measurable notion of structural "
        "correctness when research objects are created (acquisition plus first export), complementing "
        "post-hoc harmonization (e.g., ComBat) and representation standards (DICOM, BIDS).",
    )
    bullets(
        doc,
        [
            "Conceptual frame: AI-Ready Acquisition Schema (AI-RAS), acquisition contracts, protocol versioning, "
            "raw vs derived discipline, export gates, manifests.",
            "Operational artifact: modality-level SOP-style guidance (AI_Ready_Acquisition_Standards_by_Modality.docx) "
            "suitable as Methods supplement material.",
        ],
    )

    h(doc, "2. CIDUR_data: what was characterized", level=1)
    p(
        doc,
        "Exploratory analysis of /CIDUR_data (XNAT-style layout: EP*/EP*/*_MR_*/scans/.../DICOM/files/*.dcm) "
        "yielded the following empirical picture.",
    )
    h(doc, "2.1 Scanner hardware clusters", level=2)
    bullets(
        doc,
        [
            "GE Medical Systems | SIGNA Premier | 3 T — largest share of MR sessions.",
            "Siemens | Skyra | 3 T.",
            "Siemens Healthineers | MAGNETOM Vida Fit | 3 T.",
            "GE Medical Systems | SIGNA Artist | 1.5 T — smaller subset.",
            "InstitutionName, StationName, DeviceSerialNumber were empty in sampled DICOM (likely de-identification); "
            "site identity is inferred from hardware/software clusters, not explicit tags.",
        ],
    )
    h(doc, "2.2 Naming and structure heterogeneity", level=2)
    bullets(
        doc,
        [
            "Scan folder names (e.g., N-DESCRIPTION) differ strongly across vendors; little overlap between GE Premier "
            "and Siemens Skyra; limited overlap even between Skyra and Vida Fit.",
            "GE Premier and GE Artist share more naming patterns (same vendor family).",
            "SeriesDescription often does not match folder strings; some GE series have empty SeriesDescription.",
            "Diffusion: GE-oriented naming (e.g., Ax_DTI_50directions, DWI MUSE) vs Siemens 64-direction AP/PA "
            "with separate derivative series (ADC, FA, TRACEW).",
            "Transfer syntax: mostly JPEG lossless; subset of GE Premier also explicit VR little endian — automation "
            "must tolerate mixed encoding profiles.",
        ],
    )
    h(doc, "2.3 Implications for data cleaning", level=2)
    bullets(
        doc,
        [
            "No single global regex for modality; per-vendor or per-scanner-cluster mapping tables are required.",
            "High risk of confusing raw DWI with derived maps if naming and manifest rules are weak.",
            "Phase-encoding and readout metadata must be present for distortion pipelines; missing tags increase "
            "manual QC and failure rates.",
            "Human wall-clock cleaning time is not stored in DICOM; only measurable with prospective logging or "
            "proxy metrics (validator failures, manual fix counts, pipeline warnings).",
        ],
    )
    h(doc, "2.4 Quantitative feasibility from CIDUR alone", level=2)
    bullets(
        doc,
        [
            "Suitable for: metadata completeness rates, naming diversity metrics, scanner-stratified summaries, "
            "rule-based DBI scores, automated conversion pass/fail if a fixed tool is run and logged.",
            "Heuristic modality counts suggested substantial series-level material for structural, DWI, FLAIR, SWI, "
            "ASL; BOLD and fieldmaps sparser (fieldmaps more common on Vida in one pass) — case studies should "
            "report N honestly.",
            "Not available from images alone: true time-to-clean unless external logs are added.",
        ],
    )

    h(doc, "3. External standards (role in the argument)", level=1)
    h(doc, "3.1 Acquisition vs post-acquisition", level=2)
    bullets(
        doc,
        [
            "DICOM: normative representation at acquisition/exchange; does not standardize one research protocol.",
            "BIDS: file naming and sidecars after conversion; indirectly pressures acquisition/export to preserve "
            "parameters (e.g., PhaseEncodingDirection, TotalReadoutTime) that pipelines expect.",
            "BEP033 / aDWI-BIDS: advanced diffusion sharing; supports modality-deep case study citations.",
            "FAIR: metadata planning early; not modality-specific.",
            "Harmonization literature (e.g., ComBat, Fortin et al. on multi-site DTI): motivates correcting batch "
            "effects after the fact — the proposed work targets reducing structural failure modes at birth.",
        ],
    )

    h(doc, "4. NeuroImage targeting", level=1)
    p(
        doc,
        "NeuroImage accepts original research, methods papers, databases, and theory/conceptual contributions. "
        "A methods-oriented manuscript with quantitative audit + operationalized DBI + dcm2niix batch outcomes "
        "is a plausible fit. Pure standards text without measurement is weaker unless framed as theory "
        "with strong quantitative illustration.",
    )
    bullets(
        doc,
        [
            "Follow the current Guide for authors (highlights, graphical abstract optional, ethics, open access APC).",
            "Open science: public validator code and clear algorithms strengthen reproducibility claims.",
        ],
    )

    h(doc, "5. Strong claims and how to prove them", level=1)
    h(doc, "5.1 Primary claims (candidates)", level=2)
    bullets(
        doc,
        [
            "Problem: multi-vendor data often fails automation at the acquisition–analysis boundary (metadata, "
            "series identity, raw vs derived).",
            "Concept: quality scored at data birth (DBI), not only after harmonization.",
            "Prescriptive: small set of acquisition/export conventions reduces structural failure modes.",
            "Operational: enforcing conventions improves measurable outcomes (validator pass rate, conversion success, "
            "optional downstream stability).",
        ],
    )
    h(doc, "5.2 Evidence (without requiring cleaning time)", level=2)
    bullets(
        doc,
        [
            "Structural audit on full cohort: tag tables, empty description rates, diffusion ambiguity counts, "
            "by scanner cluster.",
            "Implemented DBI v1 + results table and figure.",
            "dcm2niix batch log (pinned version) — warnings/errors per series or session; pass/fail rubric → Table 2.",
            "Further pipelines (e.g., FSL/MRtrix, ML) remain out of scope for this manuscript unless added later.",
            "Proxies for effort if desired: manual relabel count, exception log entries, inter-rater series-role "
            "disagreement — not the same as hours, but objective.",
        ],
    )
    h(doc, "5.3 Cleaning time", level=2)
    p(
        doc,
        "Wall-clock cleaning time is not required for a credible NeuroImage methods paper if other objective "
        "automation-facing metrics are reported. Prospective timing strengthens an operations narrative but is optional.",
    )
    h(doc, "5.4 Locked outcomes (team decisions)", level=2)
    bullets(
        doc,
        [
            "Sharing: both analysis code and aggregate (and/or per-session) statistics intended for public release; "
            "validator packaged similarly in spirit to bids-validator (CLI, spec, tests, releases).",
            "Cohort: CIDUR_data is the paper cohort; identifiers will be de-identified to BIDS-style subject labels "
            "(e.g., sub-0001, sub-0002, …). Re-count N after exclusions in the de-ID pipeline.",
            "Primary paper outcomes (v1, both required): (1) DBI v1 scores and audit; (2) dcm2niix batch success "
            "(pinned version, pass/fail or warning policy defined in Methods). No other conversion tools in scope for now.",
            "Future work (explicitly out of scope for v1 manuscript): end-to-end BIDS layout (e.g., HeuDiConv / "
            "dcm2bids with versioned heuristics) and correlation of DBI with bids-validator severity. Discussion "
            "will frame DBI as intended to ease future BIDS conversion, without reporting full BIDS outcomes yet.",
        ],
    )

    h(doc, "6. Paper structure suggestion", level=1)
    bullets(
        doc,
        [
            "General multi-modality conventions (cross-cutting): protocol ID, Tier A/B metadata policy, export profile, "
            "manifest, raw vs derived.",
            "Per-modality case studies (authorship split possible): shared validator spec and outcome definitions so "
            "the paper reads as one study.",
            "Primary quantitative backbone: DBI + scanner-stratified audit + dcm2niix ingest outcomes (Table 2).",
        ],
    )

    h(doc, "7. Do we have everything needed?", level=1)
    bullets(
        doc,
        [
            "Data sufficiency: CIDUR_data is adequate to compute most proposed quantitative items once rules and "
            "scripts are implemented and run.",
            "Study completeness: not yet — requires executed analysis, figures/tables, manuscript, IRB/data-use text, "
            "and reproducibility packaging.",
            "Assistant support: planning, DBI formalization, Python audit/conversion tooling, draft sections; "
            "institutional submission, ethics, and final approval remain with the research team.",
        ],
    )

    h(doc, "8. Credible reference papers (curated)", level=1)
    p(
        doc,
        "Abbreviated list for Introduction/Discussion; verify formatting against journal style before submission.",
    )
    bullets(
        doc,
        [
            "Gorgolewski et al. (2016). BIDS. Scientific Data. doi:10.1038/sdata.2016.44",
            "Poldrack et al. (2017). Reproducible neuroimaging. Nature Reviews Neuroscience. doi:10.1038/nrn.2016.167",
            "Norgaard et al. (2023). BIDS experiences. Nature Reviews Neuroscience. doi:10.1038/s41583-023-00762-1",
            "Fortin et al. (2017). Multi-site DTI harmonization. NeuroImage. doi:10.1016/j.neuroimage.2017.08.047",
            "Johnson et al. (2007). ComBat (genomics origin). Biostatistics.",
            "Esteban et al. (2017). MRIQC. PLOS ONE. doi:10.1371/journal.pone.0184661",
            "Esteban et al. (2019). Crowdsourced MRI QC data. Scientific Data. doi:10.1038/s41597-019-0035-4",
            "Botvinik-Nezer et al. (2020). NARPS / analyst variability. Nature. doi:10.1038/s41586-020-2314-9",
            "Van Essen et al. (2013). HCP overview. NeuroImage. doi:10.1016/j.neuroimage.2013.05.041",
            "DICOM Standard (NEMA). https://www.dicomstandard.org/current/",
            "aDWI-BIDS / BEP033 lineage. arXiv:2103.14485; bids.neuroimaging.io extensions.",
            "Add 2–4 outcome-specific citations (converter, ML domain shift, etc.) once methods are fixed.",
        ],
    )

    h(doc, "9. Next concrete milestones", level=1)
    numbered = [
        "Lock 1–3 primary claims and target figure list.",
        "Freeze DBI v1 operational definition (components, weights or binary rules).",
        "Run audit on all MR sessions → Table 1 + Figure 1 by scanner cluster.",
        "dcm2niix batch log → Table 2 (required). BIDS validator / full BIDS recipe → future work only.",
        "Draft manuscript + supplement; internal review; NeuroImage submission.",
    ]
    for item in numbered:
        doc.add_paragraph(item, style="List Number")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.add_run("Note: ").bold = True
    foot.add_run(
        "This file is a synthesis of collaborative planning dialogue, not a peer-reviewed document. "
        "Empirical counts (e.g., exact N per modality) should be recomputed from the frozen cohort and code version "
        "cited in the final paper."
    )

    out = "/mnt/nfs/home/urmc-sh.rochester.edu/pndagiji/Documents/vecta/thought_experiment.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
