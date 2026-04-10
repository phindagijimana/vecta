#!/usr/bin/env python3
"""
Generate Phase 0 sign-off + Phase 3 frozen-results appendix (NeuroImage methods paper).

Reads (if present):
  cidur_dbi/outputs/run_metadata.json
  cidur_dbi/outputs_dcm2niix_v1freeze/dcm2niix_environment.json
  cidur_dbi/outputs_dcm2niix_v1freeze/conversion_log.csv
  cidur_dbi/outputs_dcm2niix_v1freeze/table2_conversion_by_scanner.csv

Outputs:
  vecta/phase0_signoff_neuroimage.docx
  vecta/phase3_frozen_for_manuscript.docx  (partial if Phase 3 still running)
"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
CIDUR_DBI = ROOT / "cidur_dbi"
OUT_AUDIT = CIDUR_DBI / "outputs"
OUT_P3 = CIDUR_DBI / "outputs_dcm2niix_v1freeze"


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def load_json(path: Path) -> dict | None:
    if not path.is_file():
        return None
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def build_phase0_doc(meta: dict | None) -> Document:
    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Phase 0 — Sign-off sheet (NeuroImage methods manuscript)")
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph()

    n_sess = meta.get("n_mr_sessions", "[TEAM: from audit]") if meta else "[run_audit.py → run_metadata.json]"
    n_ser = meta.get("n_series_rows", "—") if meta else "—"
    mean_dbi_s = meta.get("mean_dbi_series", "—") if meta else "—"
    mean_dbi_sess = meta.get("mean_dbi_session", "—") if meta else "—"
    audit_root = meta.get("root", "[path]") if meta else "[path]"

    p(
        doc,
        "Dataset snapshot (from latest DBI audit; cite the frozen run you use in the paper):",
    )
    bullets(
        doc,
        [
            f"MR sessions (N): {n_sess}",
            f"Series / scan folders (N): {n_ser}",
            f"Mean DBI (series-level): {mean_dbi_s}",
            f"Mean DBI (session mean, localizers excluded as implemented): {mean_dbi_sess}",
            f"Audit root (resolved): {audit_root}",
        ],
    )

    h(doc, "1. Article type", level=2)
    bullets(
        doc,
        [
            "Proposed: Methods-primary NeuroImage paper: DBI v1 operationalization + retrospective cohort audit + dcm2niix ingest outcomes.",
            "[ ] Confirmed by team    [ ] Alternative: _________________________",
        ],
    )

    h(doc, "2. Primary claims (1–3)", level=2)
    bullets(
        doc,
        [
            "P1. DBI v1 is a reproducible 0–1 summary of automation-oriented structural readiness of DICOM series (not clinical image quality), with modality-aware components M/P/G/S/N per frozen specification.",
            f"P2. In this cohort (N = {n_sess} MR sessions), DBI varies across scanner clusters; Table 1 / Figure 1 report scanner-stratified summaries.",
            "P3. Batch dcm2niix conversion under a pre-specified pass rule yields cohort-wide ingest outcomes by scanner and heuristic series class (Table 2 / Figure 2).",
        ],
    )

    h(doc, "3. Secondary claims", level=2)
    bullets(
        doc,
        [
            "S1. Subscores are auditable against explicit DICOM tags and YAML rules.",
            "S2. Series class labels are rule-based (folder + SeriesDescription); heuristic unless Phase 4 human validation is completed and claimed.",
            "S3. Full BIDS layout and bids-validator correlation are explicit future work (not v1 outcomes).",
            "S4. dcm2niix version and CLI are pinned in Methods and in dcm2niix_environment.json.",
            "[TEAM: add S5 if needed]",
        ],
    )

    h(doc, "4. Figure and table plan (v1 locked outcomes)", level=2)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "ID"
    hdr[1].text = "Source file (repo)"
    hdr[2].text = "Purpose"
    rows = [
        ("Table 1", "cidur_dbi/outputs/table1_dbi_by_scanner.csv", "DBI by scanner cluster (session-level aggregates)"),
        ("Figure 1", "cidur_dbi/outputs/figure1_dbi_by_scanner.png", "Distribution of session-mean DBI by scanner"),
        ("Table 2", "outputs_dcm2niix_v1freeze/table2_conversion_by_scanner_class.csv", "dcm2niix pass rate by scanner × heuristic class"),
        ("Figure 2", "outputs_dcm2niix_v1freeze/figure2_dcm2niix_pass_rate_heatmap.png", "Heatmap of mean conversion pass rate"),
        ("Supp. (opt.)", "cidur_dbi/outputs/figure_supp_dbi_by_class.png", "Mean DBI by heuristic series class"),
    ]
    for a, b, c in rows:
        cells = tbl.add_row().cells
        cells[0].text = a
        cells[1].text = b
        cells[2].text = c

    h(doc, "5. IRB / data sharing (team completes)", level=2)
    bullets(
        doc,
        [
            "Allowed cohort description in manuscript: [TEAM]",
            "Identifiers in public materials: [TEAM: none / aggregates only / …]",
            "Code sharing tier: [ ] Public repo + DOI   [ ] On request   [ ] Institutional only",
            "Data sharing tier: [ ] Not shared   [ ] De-ID subset   [ ] Exemplar DICOM only",
            "Reproducibility statement in paper must match the tier selected above (Phase 6).",
        ],
    )

    h(doc, "6. De-identification stage for DBI (pick one; Methods text)", level=2)
    bullets(
        doc,
        [
            "[ ] Option A: DBI computed on research export DICOM as ingested; public materials = aggregates + de-ID exemplars only.",
            "[ ] Option B: DBI computed on de-identified DICOM; document tags blanked by de-ID tool.",
        ],
    )

    h(doc, "7. Optional manuscript elements", level=2)
    bullets(
        doc,
        [
            "Phase 2 QC: spot-check k series (qc_spotcheck_template.csv) — [ ] Done  [ ] Not used; one-line limitation in Discussion.",
            "Phase 4: human labels + κ / confusion matrix — [ ] Claimed (supplement)  [ ] Not claimed; classification described as heuristic.",
        ],
    )

    h(doc, "8. Author sign-off", level=2)
    p(doc, "We agree Phase 0 is locked for v1.0 internal circulation (subject to final Phase 3 numbers).")
    bullets(
        doc,
        [
            "Corresponding author: __________________  Date: ________",
            "Senior author / PI: _____________________  Date: ________",
            "Imaging lead: ___________________________  Date: ________",
        ],
    )

    return doc


def build_phase3_appendix(p3_env: dict | None, conv_csv: Path | None) -> Document:
    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Phase 3 — Frozen outputs for manuscript (dcm2niix)")
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph()

    if p3_env is None or conv_csv is None or not conv_csv.is_file():
        p(
            doc,
            "Phase 3 full cohort run not found or still in progress. When complete, re-run:",
        )
        p(
            doc,
            "  python3 build_phase5_prereq_docx.py",
        )
        p(
            doc,
            "Expected directory: cidur_dbi/outputs_dcm2niix_v1freeze/ "
            "(conversion_log.csv, dcm2niix_environment.json, table2_*.csv, figure2_*.png).",
        )
        p(
            doc,
            "Methods placeholder until frozen:",
        )
        p(
            doc,
            "DICOM series were converted with dcm2niix [VERSION from dcm2niix -v on analysis machine] "
            "using the command pattern dcm2niix -z y -f '%n_%s' -o <output_dir> <dicom_input_dir>, "
            "with one output directory per XNAT scan folder (DICOM in …/resources/DICOM/files). "
            "Conversion success was defined as exit code 0 and at least one .nii or .nii.gz in that output directory. "
            "Outcomes were summarized by scanner cluster and heuristic series class (Table 2; Figure 2).",
        )
        return doc

    df = pd.read_csv(conv_csv)
    ran = df[df["status"].isin(("ran", "backfill_from_nifti"))]
    skip = df[df["status"] == "skip_no_dicom"]
    n_ran = len(ran)
    n_skip = len(skip)
    if n_ran == 0:
        p(doc, "conversion_log.csv has no status==ran rows; check run.")
        return doc

    pass_n = int(ran["convert_pass"].apply(lambda x: str(x).lower() in ("true", "1")).sum())
    rate = pass_n / n_ran if n_ran else 0.0
    n_backfill = int((df["status"] == "backfill_from_nifti").sum())
    ver = p3_env.get("dcm2niix_version", "?")
    run_utc = p3_env.get("run_utc", "?")
    rubric = p3_env.get("rubric", "")

    h(doc, "Pinned tooling", level=2)
    bullets(
        doc,
        [
            f"dcm2niix version (from batch run): {ver}",
            f"Run timestamp (UTC): {run_utc}",
            f"CLI pattern: {p3_env.get('cli_pattern', '')}",
            f"Rubric: {rubric}",
        ],
    )

    h(doc, "Cohort counts (conversion log)", level=2)
    count_bullets = [
        f"Series in Table 2 denominator (status=ran or backfill_from_nifti): {n_ran}",
        f"Series skipped (no DICOM in folder): {n_skip}",
        f"Passing rubric (convert_pass=True): {pass_n} ({100.0 * rate:.2f}% of denominator)",
    ]
    if n_backfill:
        count_bullets.append(
            f"Rows reconstructed from NIfTI tree only (status=backfill_from_nifti): {n_backfill} — "
            "Methods should state that pass/fail was inferred from output files, not recovered dcm2niix exit codes."
        )
    bullets(doc, count_bullets)

    tbl2_sc = OUT_P3 / "table2_conversion_by_scanner.csv"
    if tbl2_sc.is_file():
        sg = pd.read_csv(tbl2_sc)
        h(doc, "Table 2 (scanner marginal) — paste-ready summary", level=2)
        t2 = doc.add_table(rows=1, cols=4)
        t2.style = "Table Grid"
        hcells = t2.rows[0].cells
        for i, name in enumerate(["scanner_cluster", "n", "n_pass", "pass_rate"]):
            hcells[i].text = name
        for _, row in sg.head(20).iterrows():
            c = t2.add_row().cells
            c[0].text = str(row.get("scanner_cluster", ""))[:80]
            c[1].text = str(int(row.get("n", 0)))
            c[2].text = str(int(row.get("n_pass", 0)))
            c[3].text = str(row.get("pass_rate", ""))

    h(doc, "Methods paragraph (copy into manuscript)", level=2)
    cli_pat = p3_env.get("cli_pattern", 'dcm2niix -z y -f "%n_%s" -o <out> <in>')
    methods = (
        f"We converted each series’ DICOM folder (…/resources/DICOM/files) using dcm2niix "
        f"({ver}) with the command pattern {cli_pat}. "
        f"Conversion was judged successful if the process exited with code 0 and at least one NIfTI file "
        f"(.nii or .nii.gz) was written to the series output directory. "
        f"Across {n_ran} series with DICOM present, {pass_n} ({100.0 * rate:.1f}%) met this criterion. "
        f"Series without DICOM files in the expected location (n = {n_skip}) were not submitted to dcm2niix and are "
        f"excluded from scanner-by-class summaries in Table 2 and Figure 2. "
        f"Full per-series logs are available in conversion_log.csv for the frozen run dated {run_utc}."
    )
    if n_backfill:
        methods += (
            f" For {n_backfill} series, rows labeled backfill_from_nifti were reconstructed from the saved NIfTI "
            "outputs only (same success rule: ≥1 NIfTI in the series directory); original dcm2niix exit codes "
            "were not recovered."
        )
    p(doc, methods)

    h(doc, "Results sentences (draft)", level=2)
    p(
        doc,
        f"Standard DICOM-to-NIfTI conversion with dcm2niix succeeded for {100.0 * rate:.1f}% of series "
        f"with available DICOM (n = {n_ran}), with variation by scanner cluster and heuristic series class "
        f"(Table 2; Figure 2).",
    )

    h(doc, "File paths to cite in supplement / reproducibility", level=2)
    bullets(
        doc,
        [
            str((CIDUR_DBI / "outputs_dcm2niix_v1freeze" / "conversion_log.csv").resolve()),
            str((CIDUR_DBI / "outputs_dcm2niix_v1freeze" / "dcm2niix_environment.json").resolve()),
            str((CIDUR_DBI / "outputs_dcm2niix_v1freeze" / "table2_conversion_by_scanner_class.csv").resolve()),
            str((CIDUR_DBI / "outputs_dcm2niix_v1freeze" / "figure2_dcm2niix_pass_rate_heatmap.png").resolve()),
        ],
    )

    return doc


def main() -> None:
    meta = load_json(OUT_AUDIT / "run_metadata.json")
    p3_env = load_json(OUT_P3 / "dcm2niix_environment.json")
    conv_csv = OUT_P3 / "conversion_log.csv"

    p0 = build_phase0_doc(meta)
    out0 = ROOT / "phase0_signoff_neuroimage.docx"
    p0.save(out0)
    print(out0)

    p3 = build_phase3_appendix(p3_env, conv_csv)
    out3 = ROOT / "phase3_frozen_for_manuscript.docx"
    p3.save(out3)
    print(out3)


if __name__ == "__main__":
    main()
