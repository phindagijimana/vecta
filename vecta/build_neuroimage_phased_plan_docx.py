#!/usr/bin/env python3
"""Generate NeuroImage paper phased planning document (DOCX)."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text):
    doc.add_paragraph(text)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def phase_table(doc, rows):
    """rows: list of (Activity, Owner, Deliverable)"""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Activity"
    hdr[1].text = "Owner"
    hdr[2].text = "Deliverable / output"
    for a, o, d in rows:
        c = tbl.add_row().cells
        c[0].text = a
        c[1].text = o
        c[2].text = d


def main():
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("NeuroImage paper: phased research plan")
    r.bold = True
    r.font.size = Pt(18)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "CIDUR-based methods manuscript — AI/automation-inspired acquisition standards & DBI"
    ).italic = True
    doc.add_paragraph()

    p(
        doc,
        "This document breaks the remaining work into phases with activities, owners, dependencies, "
        "and concrete deliverables. “AI-assisted” means collaborative work with coding and drafting support; "
        "“Research team” means institutional accountability, ethics, execution on restricted systems, and "
        "final scientific approval.",
    )

    h(doc, "Executive summary", level=1)
    bullets(
        doc,
        [
            "Phases 0–2: scope, ethics constraints, DBI v1 spec, audit code, first results (Table 1 / Figure 1).",
            "Phase 3 (required for v1): dcm2niix batch proof → Table 2 / Figure 2 as planned.",
            "Phase 4 (optional): human-label validation for classification accuracy if claimed.",
            "Phases 5–7: full manuscript, supplement, reproducibility package, internal review, submission, revision.",
            "Future (not v1 paper): versioned BIDS conversion heuristics + bids-validator; correlate DBI with validator severity.",
        ],
    )

    h(doc, "Role split (reference)", level=1)
    h(doc, "AI-assisted (end-to-end support)", level=2)
    bullets(
        doc,
        [
            "Narrow NeuroImage story: 1–3 claims, article type (methods vs original), figure/table plan.",
            "Operationalize DBI: explicit rules, scoring rubric, pseudocode, Python on CIDUR_data.",
            "Audit pipeline: DICOM tree walk, tag extraction, missing-field flags, series classification rules, "
            "scanner-stratified summaries.",
            "Required for v1: batch dcm2niix (pinned version), log success/warnings, pass/fail definition → Table 2.",
            "Draft sections: Introduction framing, Methods, Results skeleton, Discussion (limits, related work).",
            "Supplement drafts: SOP excerpts, extra tables, validator specification.",
            "Reproducibility package: README, config template, rerun instructions (within policy).",
        ],
    )
    h(doc, "Research team only", level=2)
    bullets(
        doc,
        [
            "IRB / consent / data use: what may be described, shared, or placed in a public repository.",
            "Authorship order, competing interests, funding statements, final approval of all text.",
            "Running heavy or queued jobs on institutional systems when automation cannot execute here.",
            "NeuroImage portal: submission, cover letter, revisions, reviewer responses.",
            "Ground truth for ambiguous DICOM: protocol truth, manual labels for validation subsets.",
        ],
    )

    h(doc, "Prerequisites from the team (start Phase 0)", level=1)
    bullets(
        doc,
        [
            "LOCKED (update if changed): (a) Public release of both code and statistics; validator packaged in the spirit of bids-validator.",
            "(b) Cohort = CIDUR_data after de-identification to sub-0001, sub-0002, …; cite final N after exclusions.",
            "(c) Primary outcomes = DBI audit + dcm2niix batch proof (both required for v1; pin dcm2niix version in Methods).",
            "(d) Full BIDS layout + bids-validator vs DBI → explicit future work (not required for v1 submission).",
            "(e) Optional: 20–50 series folders with human-assigned modality/role labels for classification accuracy.",
            "(f) Optional: prospective cleaning-time logs — not required for core claims.",
        ],
    )

    h(doc, "Data and tooling assumptions", level=1)
    bullets(
        doc,
        [
            "Primary imaging root: CIDUR_data (XNAT-style EP*/EP*/*_MR_*/scans/.../DICOM).",
            "Supporting artifacts: vecta/ (DOCX generators, phased plan, future audit/DBI scripts).",
            "Empirical claims in the paper must cite the frozen code version and cohort snapshot used for numbers.",
        ],
    )

    doc.add_page_break()

    # --- Phase 0
    h(doc, "Phase 0 — Scope, governance, and story lock", level=1)
    p(doc, "Goal: one-page agreement so later phases do not rework direction.")
    phase_table(
        doc,
        [
            (
                "Confirm cohort N and IRB-permitted descriptions",
                "Research team",
                "Email memo or protocol excerpt: allowed claims, sharing tier",
            ),
            (
                "Choose article type (methods-primary vs original with methodological core)",
                "Team + AI-assisted",
                "Half-page decision note",
            ),
            (
                "Write 1–3 primary claims + 3–5 secondary claims",
                "Team + AI-assisted",
                "Claims sheet (bulleted)",
            ),
            (
                "Figure/table plan (draft list with purpose per panel)",
                "AI-assisted + team sign-off",
                "Figure plan document (e.g., Fig1 audit by scanner, Fig2 DBI components)",
            ),
            (
                "Confirm v1 outcomes: DBI + dcm2niix only (no other converters in scope)",
                "Research team",
                "Locked outcome paragraph; optional preregistration",
            ),
        ],
    )
    p(doc, "Exit criteria: signed-off claims + figure list + sharing tier + explicit DBI + dcm2niix outcome plan.")
    p(doc, "Estimated effort (team): 0.5–2 sessions; (AI-assisted): 1 iteration on drafts.")

    h(doc, "Phase 1 — DBI v1 formal specification", level=1)
    p(doc, "Goal: reproducible definition of Data Birth Integrity for this study (not necessarily final forever).")
    phase_table(
        doc,
        [
            (
                "List DBI components (e.g., metadata completeness, naming, gradient integrity for DWI, spatial, drift)",
                "AI-assisted draft + team edit",
                "DBI v1 spec (Markdown or DOCX subsection)",
            ),
            (
                "Per component: binary pass/fail vs 0–1 score; weights if composite",
                "Team decision",
                "Scoring table in spec",
            ),
            (
                "Per-modality minimum rules (which components apply to which series types)",
                "Team + AI-assisted",
                "Modality × component matrix",
            ),
            (
                "Pseudocode for session-level and series-level aggregation",
                "AI-assisted",
                "Appendix for Methods / supplement",
            ),
        ],
    )
    p(doc, "Exit criteria: frozen DBI v1.0 document with version date and changelog stub.")
    p(doc, "Dependencies: none beyond Phase 0 lock — dcm2niix is in scope for v1 alongside DBI.")

    h(doc, "Phase 2 — Audit pipeline implementation and primary results", level=1)
    p(doc, "Goal: run on full MR session set; produce Table 1 + Figure 1 for the paper.")
    phase_table(
        doc,
        [
            (
                "Python package: walk CIDUR_data, one record per scan folder or per series sample",
                "AI-assisted; team runs if needed",
                "cidur_audit/ or vecta/cidur_dbi/ with requirements.txt",
            ),
            (
                "Extract DICOM tags for checklist; handle transfer syntax / read failures",
                "AI-assisted",
                "Logged exceptions per file",
            ),
            (
                "Assign scanner cluster (manufacturer, model, field)",
                "AI-assisted",
                "Column in per-series CSV",
            ),
            (
                "Compute DBI v1 subscores and composite",
                "AI-assisted",
                "per_series.csv, per_session.csv summaries",
            ),
            (
                "Scanner-stratified tables and plots",
                "AI-assisted + team QC",
                "Table 1 (LaTeX/Word) + Figure 1 (e.g., stacked bars or heatmap)",
            ),
            (
                "Spot-check 10–20 random series against console or second reader",
                "Research team",
                "Short QC log",
            ),
        ],
    )
    p(doc, "Exit criteria: frozen results files + figure/table sources committed with a tag or Zenodo snapshot if allowed.")
    p(doc, "Dependencies: Phase 1 frozen spec; pydicom (or equivalent) on execution environment.")

    doc.add_page_break()

    h(doc, "Phase 3 — dcm2niix batch proof (required for v1; not full BIDS)", level=1)
    p(
        doc,
        "Goal: second core empirical outcome alongside DBI — standard DICOM→NIfTI ingest stress test on the cohort. "
        "Full BIDS conversion (e.g., HeuDiConv / dcm2bids) and bids-validator correlation remain future work (Phase 8). "
        "Discussion may motivate DBI as easing eventual BIDS setup.",
    )
    phase_table(
        doc,
        [
            (
                "Pin tool versions (e.g., dcm2niix) and document CLI",
                "Team + AI-assisted",
                "Methods paragraph + environment file",
            ),
            (
                "Batch convert sample or full cohort; capture stderr/stdout and exit codes",
                "Team run + AI-assisted scripts",
                "conversion_log.csv",
            ),
            (
                "Define pass/fail (e.g., NIfTI exists, no fatal error, warning budget)",
                "Team",
                "Rubric in supplement",
            ),
            (
                "Cross-tab pass/fail × scanner cluster × modality heuristic",
                "AI-assisted",
                "Table 2 + optional Figure 2",
            ),
        ],
    )
    p(doc, "Exit criteria: Table 2 (and optional Figure 2) ready for manuscript — required for v1 submission package.")
    p(doc, "Dependencies: software installed on cluster; legal to run conversion on data.")

    h(doc, "Phase 4 — Optional: validation subset with human labels", level=1)
    p(doc, "Goal: report accuracy/precision of automated series-role classification (if claimed).")
    phase_table(
        doc,
        [
            (
                "Sample 20–50 series stratified by scanner",
                "Research team",
                "Spreadsheet: path, true_role, notes",
            ),
            (
                "Compare rule-based or ML classifier to labels",
                "AI-assisted",
                "Confusion matrix, accuracy, kappa",
            ),
            (
                "Revise rules if systematic errors found",
                "Team + AI-assisted",
                "DBI v1.1 or classification appendix",
            ),
        ],
    )
    p(doc, "Exit criteria: supplement table or main text paragraph; or phase skipped if not a claim.")

    h(doc, "Phase 5 — Manuscript and supplement", level=1)
    phase_table(
        doc,
        [
            (
                "Introduction: problem, gap vs BIDS/DICOM/harmonization, contributions",
                "AI-assisted draft + team rewrite",
                "intro.docx / .tex",
            ),
            (
                "Methods: cohort, scanner groups, DBI definition, audit software, statistics",
                "AI-assisted + team",
                "methods section",
            ),
            (
                "Results: align every claim to Figure/Table",
                "AI-assisted + team",
                "results section",
            ),
            (
                "Discussion: limitations (N, single archive), related work, future multi-site",
                "AI-assisted + team",
                "discussion section",
            ),
            (
                "Supplement: SOP excerpts from AI_Ready_Acquisition_*, full tag list, extra figures",
                "AI-assisted + team",
                "supplement.pdf",
            ),
        ],
    )
    p(doc, "Exit criteria: full draft v1.0 internally circulable.")

    h(doc, "Phase 6 — Reproducibility and data/code packaging", level=1)
    phase_table(
        doc,
        [
            (
                "README: how to install, configure paths, rerun tables (no raw data if policy forbids)",
                "AI-assisted",
                "README.md",
            ),
            (
                "License and citation file for code",
                "Research team",
                "LICENSE, CITATION.cff optional",
            ),
            (
                "Archive snapshot (Git tag, Zenodo, institutional repo) per sharing tier",
                "Research team",
                "DOI or internal accession ID",
            ),
        ],
    )
    p(doc, "Exit criteria: reproducibility statement text for manuscript matches actual availability.")

    h(doc, "Phase 7 — Internal review, submission, revision", level=1)
    phase_table(
        doc,
        [
            ("Coauthor review rounds", "All authors", "Tracked comments resolved"),
            ("Graphical abstract and highlights (NeuroImage)", "Team", "Portal-ready assets"),
            ("Cover letter emphasizing novelty and scope", "Team + AI-assisted draft", "PDF"),
            ("Submit to NeuroImage", "Corresponding author", "Submission ID"),
            ("Respond to reviewers", "Team + AI-assisted", "Point-by-point + revised draft"),
        ],
    )
    p(doc, "Exit criteria: acceptance or resubmission plan documented.")

    h(doc, "Phase 8 — Future work (post-v1 publication)", level=1)
    p(
        doc,
        "Not part of the first NeuroImage submission timeline. After a frozen de-identified cohort and stable DBI "
        "definition, pursue: (1) versioned BIDS conversion recipe (HeuDiConv / dcm2bids or equivalent) per scanner "
        "cluster; (2) run bids-validator on resulting layouts; (3) test whether higher DBI predicts fewer errors/warnings "
        "or lower manual curation burden.",
    )
    bullets(
        doc,
        [
            "Deliverable: second paper or major revision extension + open-sourced BIDS heuristics.",
            "Dependency: sub-* layout and documented mapping from XNAT-style series names to BIDS entities.",
        ],
    )

    doc.add_page_break()

    h(doc, "Suggested timeline (indicative)", level=1)
    p(
        doc,
        "Adjust to FTE and cluster access. Parallel tracks: Phase 1 can start drafting while Phase 0 finalizes "
        "if team agrees provisional outcome A.",
    )
    checklist = doc.add_table(rows=1, cols=3)
    checklist.style = "Table Grid"
    hc = checklist.rows[0].cells
    hc[0].text = "Phase"
    hc[1].text = "Weeks (indicative)"
    hc[2].text = "Notes"
    rows = [
        ("0 Scope lock", "1", "Blocking for everything else"),
        ("1 DBI v1 spec", "1–2", "Overlaps with early coding sketch"),
        ("2 Audit + Fig1/Table1", "2–4", "Largest engineering block"),
        ("3 dcm2niix proof", "1–3", "Required; can parallel writing after Phase 2"),
        ("4 Human labels", "1–2", "Optional; can sample during Phase 2 QC"),
        ("5 Manuscript v1", "2–4", "Start prose once Table 1 exists"),
        ("6 Repro package", "1", "Often during revision"),
        ("7 Submit + revise", "8–20+", "Journal-dependent"),
        ("8 BIDS + validator (future)", "TBD", "Separate milestone after v1 paper"),
    ]
    for row in rows:
        c = checklist.add_row().cells
        for i, val in enumerate(row):
            c[i].text = val

    h(doc, "Risk register (short)", level=1)
    bullets(
        doc,
        [
            "Policy change forbids public code → Plan B: aggregate-only supplement + institutional code on request.",
            "DICOM read failures on subset → Document exclusion rule and sensitivity analysis.",
            "Sparse modality (e.g., BOLD) → Frame as illustrative case study; do not overclaim N.",
            "Reviewer demands second cohort → Discuss public subset replication or limit scope in revision.",
        ],
    )

    h(doc, "Document control", level=1)
    p(doc, "Version: 1.2. v1 scope locked: DBI + dcm2niix only; BIDS/bids-validator → Phase 8. Update dates/owners in Word as needed.")

    out = "/mnt/nfs/home/urmc-sh.rochester.edu/pndagiji/Documents/vecta/neuroimage_paper_phased_plan.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
