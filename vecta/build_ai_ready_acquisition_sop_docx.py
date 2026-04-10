#!/usr/bin/env python3
"""Generate AI-ready acquisition & export SOP documentation (per modality) as .docx."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def p(doc, text, bold_lead=None):
    if bold_lead:
        para = doc.add_paragraph()
        para.add_run(bold_lead).bold = True
        para.add_run(text)
    else:
        doc.add_paragraph(text)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def checklist_table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, hd in enumerate(headers):
        tbl.rows[0].cells[i].text = hd
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def main():
    doc = Document()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AI-Ready Acquisition & Export Standards")
    r.bold = True
    r.font.size = Pt(20)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.add_run(
        "Data Birth Integrity (DBI)–Aligned Guidance for Multi-Vendor, Multi-Center Imaging"
    ).italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run("Document type: ").bold = True
    meta.add_run("Operational specification for acquisition, export, and automation handoff.\n")
    meta.add_run("Audience: ").bold = True
    meta.add_run(
        "MRI technologists, physicists, PACS/research IT, data engineers, and imaging scientists.\n"
    )
    meta.add_run("Note: ").bold = True
    meta.add_run(
        "This document complements DICOM, institutional IRB/policy, and study protocols; "
        "it does not replace them."
    )

    doc.add_page_break()

    # ---- TOC placeholder (Word can regenerate TOC; we list sections in intro)
    h(doc, "Table of contents (sections)", level=1)
    bullets(
        doc,
        [
            "1. Purpose, scope, and definitions",
            "2. Cross-cutting standards (all modalities)",
            "3. Structural / anatomical MRI",
            "4. Diffusion MRI (DWI / DTI / advanced DWI)",
            "5. Functional MRI (BOLD)",
            "6. Field mapping (B0 distortion correction)",
            "7. Susceptibility-weighted imaging (SWI / QSM-related)",
            "8. Arterial spin labeling (ASL) perfusion",
            "9. Dynamic susceptibility contrast (DSC) perfusion",
            "10. Contrast-enhanced structural MRI",
            "11. Quantitative parametric maps (T1, T2, PD, MT)",
            "12. MR angiography (TOF / contrast-enhanced MRA) — research handoff",
            "13. MR spectroscopy (SVS / CSI) — research handoff",
            "14. Cross-center governance, DBI scoring hooks, and references",
        ],
    )
    doc.add_page_break()

    # ---- 1 Purpose
    h(doc, "1. Purpose, scope, and definitions", level=1)
    p(
        doc,
        "Clinical and research MRI data are increasingly consumed by automated pipelines, "
        "quality systems, and machine learning models. Vendor defaults, site habits, and "
        "de-identification practices often remove or fragment the structural information those "
        "systems require. This document defines acquisition- and export-level standards that are "
        "motivated by automation, reproducible cleaning, and AI readiness, while remaining "
        "implementable across Siemens, GE, Philips, Canon, United Imaging, and other vendors.",
    )
    h(doc, "1.1 Definitions", level=2)
    bullets(
        doc,
        [
            "Acquisition contract: versioned rules that must be satisfied when a series is acquired "
            "(pulse sequence choice, naming tokens, minimum metadata presence on scanner console export).",
            "Export contract: rules satisfied before research data leave PACS, XNAT, or gateway systems "
            "(no forbidden stripping of tags, agreed transfer syntax profile, sidecar policy).",
            "Semantic role: the intended use of a series in downstream software (e.g., T1w, dwi, fmap, flair).",
            "DBI (Data Birth Integrity): structural correctness at creation/export—completeness, conformity, "
            "naming, spatial plausibility, and drift control versus the active protocol version.",
            "Raw vs derived: raw data are directly from reconstruction intended for primary fitting; "
            "derived maps (ADC, FA, synthetic contrasts) must not impersonate raw series in automation.",
        ],
    )
    h(doc, "1.2 What this document does not do", level=2)
    bullets(
        doc,
        [
            "It does not prescribe disease-specific diagnostic protocols or replace radiologist judgment.",
            "It does not mandate a single field strength, resolution, or shell scheme for all science—"
            "studies still choose biology-appropriate parameters.",
            "It does not guarantee regulatory clearance for any product; it is a research operations standard.",
        ],
    )

    doc.add_page_break()

    # ---- 2 Cross-cutting
    h(doc, "2. Cross-cutting standards (all modalities)", level=1)
    p(
        doc,
        "These rules apply to every modality section below unless explicitly overridden.",
    )
    h(doc, "2.1 Protocol identity and drift control", level=2)
    bullets(
        doc,
            [
            "Every series must carry a machine-readable protocol version token agreed by the study "
            "(e.g., PROTO_NEURO_V3) inside ProtocolName and/or SeriesDescription using a fixed pattern.",
            "Site code + study code may be prepended in a fixed order: SITE_STUDY_PROTOVERSION_SEQUENCE_ROLE.",
            "Any console software upgrade, gradient firmware change, or major sequence edit increments "
            "the protocol version; a phantom QC session must be acquired under the new version before "
            "human accrual resumes under that label.",
        ],
    )
    h(doc, "2.2 Controlled vocabulary for semantic roles", level=2)
    p(
        doc,
        "Automation should map each series to exactly one primary role using study-defined tokens "
        "(examples: T1w, T2w, FLAIR, PD, dwi, fmap, bold, swi, asl, perf, ceT1w). Each series "
        "description must contain exactly one such token from the study list.",
    )
    h(doc, "2.3 De-identification vs ML-critical metadata", level=2)
    bullets(
        doc,
        [
            "Tier A (preserve in DICOM or parallel JSON): voxel spacing, orientation, timing, echo times, "
            "flip angle, phase encoding, multiband/shots, sequence variant, vendor, model, field strength, "
            "coil elements (if allowed), software version, reconstruction type.",
            "Tier B (suppress per IRB): patient name, IDs in free text, physician names, accession in "
            "shared exports.",
            "De-identification scripts must be reviewed so Tier A is not blanked by overly aggressive rules.",
        ],
    )
    h(doc, "2.4 Export profile", level=2)
    bullets(
        doc,
        [
            "Primary research archive: uncompressed Little Endian Explicit VR or another study-fixed "
            "lossless transfer syntax; document the UID list in the study SOP.",
            "No site may supply only screen captures or viewer-only formats as the sole research copy.",
            "Each subject visit folder must include a manifest (CSV/JSON) listing series UID, role token, "
            "protocol version, and export checksum.",
        ],
    )
    h(doc, "2.5 Cross-cutting metadata checklist (export gate)", level=2)
    checklist_table(
        doc,
        ["Element", "Required for automation", "Typical DICOM tags (informative)"],
        [
            ("Modality", "Yes", "0008,0060"),
            ("Manufacturer / Model / Field strength", "Yes", "0008,0070; 0008,1090; 0018,0087"),
            ("Series / Study UID", "Yes", "0020,000E; 0020,000D"),
            ("Patient orientation / frame of reference", "If used in reg", "0020,0020; 0020,0037"),
            ("Pixel spacing / slice thickness / spacing", "Yes", "0028,0030; 0018,0050; 0018,0088"),
            ("Image orientation patient", "Yes for 3D reg", "0020,0037"),
            ("Image position patient (first slice)", "Often", "0020,0032"),
            ("ProtocolName / SeriesDescription", "Yes (with tokens)", "0018,1030; 0008,103E"),
        ],
    )

    doc.add_page_break()

    # ---- 3 Structural
    h(doc, "3. Structural / anatomical MRI", level=1)
    p(
        doc,
        "Typical roles: T1w (MP-RAGE, SPGR, IR-FSPGR), T2w, proton density, FLAIR. Used for segmentation, "
        "registration targets, cortical surfaces, lesion labels, and multimodal fusion.",
    )
    h(doc, "3.1 AI / automation / cleaning motivations", level=2)
    bullets(
        doc,
        [
            "Registration and deep segmentation assume consistent contrast, isotropy or known anisotropy, "
            "and full brain coverage without clipped FOV.",
            "Harmonization tools (ComBat, RISH, intensity normalization) fail silently when series are "
            "mislabeled (e.g., T2-FLAIR called T2w).",
        ],
    )
    h(doc, "3.2 Acquisition-level requirements", level=2)
    bullets(
        doc,
        [
            "Declare target resolution and FOV in the written protocol; if anisotropic, document expected "
            "reformat policy (on-console MPR vs offline).",
            "Avoid mixing multiple reconstructions with identical SeriesDescription; each reconstruction "
            "must differ by a suffix (_N4, _MPR_COR, etc.).",
            "For 3D acquisitions, prefer a single 3D slab per role; separate runs must have distinct "
            "protocol version or run index in the name.",
            "Motion mitigation: document whether navigators, BLADE, or repeat-acquire rules apply.",
        ],
    )
    h(doc, "3.3 Export contract", level=2)
    bullets(
        doc,
        [
            "One canonical T1w per timepoint for ML unless the study explicitly defines multi-echo T1; "
            "extras must be tagged T1w_ALT or excluded in manifest.",
            "Preserve rescale slope/intercept and real-world value mapping tags when present; do not "
            "window-level and re-save as primary research DICOM.",
        ],
    )
    h(doc, "3.4 Metadata required for pipelines (minimum)", level=2)
    checklist_table(
        doc,
        ["Requirement", "Rationale"],
        [
            ("EchoTime, RepetitionTime, InversionTime (if applicable)", "Contrast-aware models and QC"),
            ("FlipAngle", "T1-weighted modeling / synthetic MRI"),
            ("Receive coil information (if policy allows)", "Bias-field behavior differs by coil"),
            ("Parallel imaging factor (if encoded)", "Noise structure in DL denoising"),
            ("Dimensions consistent with NIfTI conversion", "Avoid off-by-one in sagittal/coronal reformats"),
        ],
    )
    h(doc, "3.5 Common pipeline failure modes (structural)", level=2)
    checklist_table(
        doc,
        ["Failure", "Mitigation at acquisition/export"],
        [
            ("MPR reformats saved with same name as source 3D", "Distinct SeriesDescription suffix per reconstruction"),
            ("Mixed echo T1 mislabeled as single T1w", "Token T1w_ME with echo list in manifest"),
            ("Incomplete head coverage", "FOV checks in technologist worksheet"),
            ("GDC burn-in on exported DICOM", "Research export from raw recon, not screenshot pipeline"),
        ],
    )
    h(doc, "3.6 DBI-style checks at birth", level=2)
    bullets(
        doc,
        [
            "Role token present and unique mapping in study table.",
            "Voxel size within study tolerance; missing spacing flagged.",
            "No duplicate SeriesInstanceUID; no zero-length series.",
        ],
    )

    doc.add_page_break()

    # ---- 4 Diffusion
    h(doc, "4. Diffusion MRI (DWI / DTI / advanced DWI)", level=1)
    p(
        doc,
        "Includes single- and multi-shell DWI, tensor imaging, and advanced encoding. This modality has "
        "the highest cleaning failure rate if derivatives are confused with raw data or if phase-encoding "
        "metadata are missing.",
    )
    h(doc, "4.1 AI / automation / cleaning motivations", level=2)
    bullets(
        doc,
        [
            "Eddy-current, motion, and susceptibility correction require phase-encoding direction and readout "
            "timing (or fieldmaps with known relationship to DWI).",
            "Tractography and microstructure models require correct b-values and gradient directions per volume.",
            "ML harmonization across sites requires consistent shell layout or explicit per-site mapping tables.",
        ],
    )
    h(doc, "4.2 Acquisition-level requirements", level=2)
    bullets(
        doc,
        [
            "Study must document shell scheme (b=0 policy, high b shells), nominal direction count per shell, "
            "and phase-encoding polarity scheme (single-PE with fmap vs blip-up/blip-down).",
            "Raw diffusion series must use a SeriesDescription containing the role token dwi and excluding "
            "derivative tokens (_ADC, _FA, _TRACE, _TRACEW, _ColFA, etc.).",
            "If the scanner automatically creates derivative series, they must be in separate series with "
            "mandatory derivative suffixes in SeriesDescription.",
            "For dual-PE acquisitions, both PE directions must include PE token and consistent run pairing "
            "in the manifest (e.g., dwi_AP / dwi_PA).",
            "Optional advanced: document waveform type if non-standard (e.g., oscillating gradients) for "
            "future BEP033-style sharing.",
        ],
    )
    h(doc, "4.3 Export contract", level=2)
    bullets(
        doc,
        [
            "Gradient tables must be recoverable: from standard DICOM diffusion tags and/or Enhanced MR "
            "functional groups, or exported sidecars (bval/bvec/json) generated at export with documented "
            "software version.",
            "Never ship NIfTI-only DWI without accompanying gradient tables and JSON metadata required by "
            "the study pipeline (BIDS-style fields where applicable).",
            "If DICOM is mosaic or multiframe, document the converter version used to produce analysis copies.",
        ],
    )
    h(doc, "4.4 Metadata required (minimum)", level=2)
    checklist_table(
        doc,
        ["Requirement", "Notes"],
        [
            ("Effective b-value per volume", "DICOM tags or private tags mapped in SOP"),
            ("Diffusion gradient orientation per volume", "Same"),
            ("PhaseEncodingDirection (conceptual)", "Must map to BIDS PE string after conversion"),
            ("TotalReadoutTime or equivalent", "For TOPUP/eddy; from DICOM or calculated per vendor SOP"),
            ("EchoTime", "Strong effect on FA/MD in some acquisitions"),
            ("Parallel imaging / multiband", "Affects SNR and artifact patterns"),
            ("Partial Fourier / phase partial", "Affects QC thresholds"),
        ],
    )
    h(doc, "4.5 Multi-shell and advanced diffusion", level=2)
    bullets(
        doc,
        [
            "Each shell must be identifiable in DICOM or in export sidecars (b-value per volume).",
            "If shell-interleaved acquisition is used, document ordering convention for converters.",
            "For CSD / NODDI / multi-compartment models, document minimum SNR policy and motion retry rules.",
            "Cardiac-triggered DWI: record trigger delay and rejected segments policy.",
        ],
    )
    h(doc, "4.6 Suggested study-level numeric guardrails (example—not universal)", level=2)
    p(
        doc,
        "Studies may adopt ranges such as: include b≈0 s/mm² volumes; at least one high-b shell commonly "
        "used in the field; sufficient directions for the chosen tensor model; near-isotropic voxels if "
        "tractography is primary. Encode chosen numbers in the acquisition contract, not in this template.",
    )
    h(doc, "4.7 Common pipeline failure modes (diffusion)", level=2)
    checklist_table(
        doc,
        ["Failure", "Mitigation"],
        [
            ("ADC map selected as input to tractography", "Derivative suffix enforcement + manifest QA"),
            ("Missing PE or readout → eddy/TOPUP crash or silent misalignment", "Export gate checklist §4.4"),
            ("Mosaic DICOM mis-converted", "Documented converter + regression test on phantom"),
            ("Private tag gradient table lost on anonymize", "Tier A preservation + sidecar at export"),
            ("Blip-up/down run pairing broken across sites", "Explicit manifest columns for paired series UIDs"),
        ],
    )
    h(doc, "4.8 DBI-style checks at birth", level=2)
    bullets(
        doc,
        [
            "Count of volumes matches expected (within tolerance for skipped volumes policy).",
            "No derivative series mapped as dwi in manifest.",
            "b-table monotonic consistency checks (no NaN directions).",
            "PE metadata present if distortion correction is in scope.",
        ],
    )

    doc.add_page_break()

    # ---- 5 fMRI
    h(doc, "5. Functional MRI (BOLD)", level=1)
    p(
        doc,
        "Resting-state and task fMRI for connectivity, GLM, and ML phenotyping.",
    )
    h(doc, "5.1 AI / automation motivations", level=2)
    bullets(
        doc,
        [
            "Slice timing correction and STC-aware models need slice order and TR.",
            "CompCor and nuisance regression need anatomical masks aligned in the same space as BOLD.",
            "Multiband sequences require accurate multiband factor and slice timing metadata.",
        ],
    )
    h(doc, "5.2 Acquisition-level requirements", level=2)
    bullets(
        doc,
        [
            "SeriesDescription must include role token bold and task or rest label (e.g., bold_rest, "
            "bold_lang).",
            "Document multiband factor, slice acceleration, partial Fourier, echo spacing or readout "
            "parameters used by distortion correction.",
            "If simultaneous multi-slice (SMS) interleaved, document interleave pattern if non-default.",
            "For task fMRI, stimulus timing files are out of DICOM scope but must be deposited with the "
            "same visit ID in the study repository.",
        ],
    )
    h(doc, "5.3 Export contract", level=2)
    bullets(
        doc,
        [
            "Preserve temporal framing: number of dynamics, discard volumes policy, and whether dummy scans "
            "were acquired.",
            "If DICOM stores one file per volume vs multiframe, document export behavior to avoid duplicate "
            "time points.",
        ],
    )
    h(doc, "5.4 Metadata required (minimum)", level=2)
    checklist_table(
        doc,
        ["Requirement", "Rationale"],
        [
            ("RepetitionTime", "STC, filtering, design matrix"),
            ("EchoTime", "Contrast, susceptibility weighting"),
            ("FlipAngle", "Signal model"),
            ("SliceTiming or explicit slice order", "STC and some DL models"),
            ("PhaseEncodingDirection + readout timing", "Distortion correction"),
            ("MultibandAccelerationFactor (if used)", "SMS timing"),
        ],
    )
    h(doc, "5.5 Common pipeline failure modes (BOLD)", level=2)
    checklist_table(
        doc,
        ["Failure", "Mitigation"],
        [
            ("Incorrect slice timing assumed", "Preserve SliceTiming or print sequence PDF to study repo"),
            ("Multiband factor missing in JSON", "Tier A tag mapping per vendor"),
            ("Duplicate dynamics after DICOM merge", "One-file-per-volume vs multiframe policy documented"),
            ("Task label wrong in SeriesDescription", "Controlled vocabulary + QC script on manifest"),
        ],
    )
    h(doc, "5.6 DBI-style checks at birth", level=2)
    bullets(
        doc,
        [
            "TR constant across volumes within tolerance.",
            "Expected number of time points vs protocol.",
            "No mixing of two tasks in one series without manifest entry.",
        ],
    )

    doc.add_page_break()

    # ---- 6 Fieldmap
    h(doc, "6. Field mapping (B0 distortion correction)", level=1)
    p(
        doc,
        "GRE fieldmaps, dual-echo phase difference, or reversed-PE references for TOPUP. Critical for "
        "fMRI and DWI when EPI distortion is non-negligible.",
    )
    h(doc, "6.1 Acquisition-level requirements", level=2)
    bullets(
        doc,
        [
            "Role token fmap (or study-defined subtypes fmap_magnitude, fmap_phase, fmap_epi).",
            "Document relationship to matched bold or dwi series (same shim, same prescription if required).",
            "For dual-echo GRE, specify echo times used in phase-difference calculation.",
        ],
    )
    h(doc, "6.2 Export contract", level=2)
    bullets(
        doc,
        [
            "Export magnitude and phase (or real/imaginary) as required by the study pipeline; do not drop "
            "one component.",
            "Manifest must link fmap series UIDs to the bold/dwi series they correct.",
        ],
    )
    h(doc, "6.3 Metadata required", level=2)
    bullets(
        doc,
        [
            "EchoTime(s), intended for fieldmap",
            "Shim settings if encoded",
            "Consistent orientation relative to matched EPI",
        ],
    )
    h(doc, "6.4 Failure modes", level=2)
    bullets(
        doc,
        [
            "fmap acquired with different shimming prescription than BOLD/DWI → ineffective correction; "
            "SOP should require paired acquisition order.",
            "Phase scaled or filtered inconsistently across sites → TOPUP failures; document vendor scaling.",
        ],
    )

    doc.add_page_break()

    # ---- 7 SWI
    h(doc, "7. Susceptibility-weighted imaging (SWI / QSM-related)", level=1)
    p(
        doc,
        "Magnitude, phase, SWI, minimum intensity projection (mIP) variants. Used in hemorrhage, iron, "
        "and some DL vessel models.",
    )
    h(doc, "7.1 Acquisition-level requirements", level=2)
    bullets(
        doc,
        [
            "Token swi with suffix for each product type (_MAG, _PHA, _SWI, _mIP).",
            "Do not reuse identical SeriesDescription for different reconstructions.",
        ],
    )
    h(doc, "7.2 Export contract", level=2)
    bullets(
        doc,
        [
            "If phase is needed for QSM, phase must not be discarded; document scaling (radians vs arbitrary).",
        ],
    )
    h(doc, "7.3 Metadata required", level=2)
    bullets(
        doc,
        [
            "EchoTime(s), flip angle, filter mode if vendor-specific",
            "Resolution and orientation",
        ],
    )

    doc.add_page_break()

    # ---- 8 ASL
    h(doc, "8. Arterial spin labeling (ASL) perfusion", level=1)
    p(
        doc,
        "CBF mapping for neurodegeneration, cerebrovascular disease, and some ML biomarkers.",
    )
    h(doc, "8.1 AI / automation motivations", level=2)
    bullets(
        doc,
        [
            "Quantification requires labeled/control pairs, post-label delay, background suppression scheme, "
            "and M0 reference policy.",
            "Multi-PLD and multi-TI schemes need explicit labeling in manifest.",
        ],
    )
    h(doc, "8.2 Acquisition-level requirements", level=2)
    bullets(
        doc,
        [
            "Token asl; include vendor sequence family in description (e.g., PCASL, PASL, pCASL).",
            "Separate M0 / proton-density reference series if acquired; tag asl_M0.",
            "Document background suppression pulses and whether a separate T1w is required for calibration.",
        ],
    )
    h(doc, "8.3 Export contract", level=2)
    bullets(
        doc,
        [
            "Do not average label/control on-console for research primary if the pipeline needs raw pairs.",
            "Preserve all dynamics needed for kinetic or M0 scaling.",
        ],
    )
    h(doc, "8.4 Metadata required (minimum)", level=2)
    checklist_table(
        doc,
        ["Requirement", "Rationale"],
        [
            ("Labeling duration and post-label delay(s)", "Quantification"),
            ("Background suppression flags", "Model choice"),
            ("M0 strategy", "Scaling"),
            ("Parallel imaging and readout type", "Artifact QC"),
        ],
    )

    doc.add_page_break()

    # ---- 9 DSC
    h(doc, "9. Dynamic susceptibility contrast (DSC) perfusion", level=1)
    p(
        doc,
        "Contrast bolus perfusion; highly sensitive to acquisition timing and leakage correction choices.",
    )
    h(doc, "9.1 Acquisition-level requirements", level=2)
    bullets(
        doc,
        [
            "Token perf_dsc; document pre-bolus baseline volumes and temporal resolution.",
            "If dual-echo DSC is used for leakage correction, both echoes must be exported with linked UIDs.",
        ],
    )
    h(doc, "9.2 Export contract", level=2)
    bullets(
        doc,
        [
            "Retain full temporal series; no vendor-only cine export as sole data.",
            "Document contrast agent type, dose, and injection time relative to scan start (Tier B linkage "
            "may be restricted—study policy decides storage).",
        ],
    )
    h(doc, "9.3 Metadata required", level=2)
    bullets(
        doc,
        [
            "TR, TE per echo, flip angle",
            "Number of dynamics, slice order",
        ],
    )

    doc.add_page_break()

    # ---- 10 CE
    h(doc, "10. Contrast-enhanced structural MRI", level=1)
    p(
        doc,
        "Post-contrast T1w (and sometimes FLAIR) for lesion enhancement. Governance overlaps safety, "
        "timing, and research reproducibility.",
    )
    h(doc, "10.1 Acquisition-level requirements", level=2)
    bullets(
        doc,
        [
            "Token ceT1w (or ceFLAIR); pre-contrast companion series must exist if subtraction or "
            "registration is planned, with matching geometry or documented reslice policy.",
            "Document delay from injection to start of acquisition for reproducible enhancement curves.",
        ],
    )
    h(doc, "10.2 Export contract", level=2)
    bullets(
        doc,
        [
            "Keep pre and post as separate series with distinct UIDs; never overwrite pre with post.",
            "IRB-compliant handling of contrast documentation in research manifests.",
        ],
    )
    h(doc, "10.3 Metadata required", level=2)
    bullets(
        doc,
        [
            "Same structural metadata as Section 3, plus timing relative to injection if recorded",
        ],
    )

    doc.add_page_break()

    # ---- 11 Quantitative maps
    h(doc, "11. Quantitative parametric maps (T1, T2, PD, MT)", level=1)
    p(
        doc,
        "Mapping sequences produce parametric images (T1 map, T2 map, PD, MTR) used in biophysical models, "
        "synthetic MRI, and some deep learning harmonization pipelines.",
    )
    h(doc, "11.1 Acquisition-level requirements", level=2)
    bullets(
        doc,
        [
            "Use tokens qmap_T1, qmap_T2, qmap_PD, qmap_MTR (or study-defined equivalents).",
            "Document flip angle train, echo train, spacing of TI or TE samples, and fitting method on-console "
            "vs offline.",
            "If magnitude and phase are both required for fitting, export both with distinct suffixes.",
        ],
    )
    h(doc, "11.2 Export contract", level=2)
    bullets(
        doc,
        [
            "Preserve raw component series if the fitting pipeline is offline; do not export only color maps.",
            "Document units (ms, s) in manifest for each map series.",
        ],
    )
    h(doc, "11.3 Metadata required", level=2)
    checklist_table(
        doc,
        ["Requirement", "Rationale"],
        [
            ("Full list of TIs or TEs used", "Reproducible fitting"),
            ("B1+ correction flag if used", "Bias in T1 estimation"),
            ("Parallel imaging factors", "SNR and bias"),
        ],
    )

    doc.add_page_break()

    # ---- 12 MRA
    h(doc, "12. MR angiography (TOF / contrast-enhanced MRA) — research handoff", level=1)
    p(
        doc,
        "Vascular imaging for stroke, aneurysm, and vessel segmentation models. Often combined with "
        "structural scans for multimodal AI.",
    )
    h(doc, "12.1 Acquisition-level requirements", level=2)
    bullets(
        doc,
        [
            "Token mra_TOF or mra_CE; include slab count and overlap if multi-slab TOF.",
            "For CE-MRA, timing relative to contrast and arterial vs venous phase intent must be documented.",
        ],
    )
    h(doc, "12.2 Export contract", level=2)
    bullets(
        doc,
        [
            "Preserve source partitions if MIP-only series are created; research primary should include "
            "non-MIP source unless study waives.",
            "Vessel segmentation models need isotropic or documented anisotropic spacing.",
        ],
    )
    h(doc, "12.3 Metadata required", level=2)
    bullets(
        doc,
        [
            "Flip angle, TR, TE, flow compensation flags if encoded",
            "Contrast timing for CE-MRA per IRB policy",
        ],
    )

    doc.add_page_break()

    # ---- 13 MRS
    h(doc, "13. MR spectroscopy (SVS / CSI) — research handoff", level=1)
    p(
        doc,
        "Single-voxel (SVS) and chemical shift imaging (CSI) produce spectra for metabolite quantification "
        "and emerging ML classifiers.",
    )
    h(doc, "13.1 Acquisition-level requirements", level=2)
    bullets(
        doc,
        [
            "Token mrs_SVS or mrs_CSI; specify nucleus (1H), TE, water suppression method, and voxel size.",
            "Document voxel placement reference series (e.g., T1w series UID used for prescription).",
        ],
    )
    h(doc, "13.2 Export contract", level=2)
    bullets(
        doc,
        [
            "Export raw FIDs or vendor spectroscopy objects per pipeline; document software used for "
            "preprocessing.",
            "Include basis set / editing scheme in study repo if applicable (e.g., MEGA, HERMES).",
        ],
    )
    h(doc, "13.3 Metadata required", level=2)
    bullets(
        doc,
        [
            "Spectral width, number of points, averaging count",
            "Voxel position and orientation relative to anatomical series",
        ],
    )

    doc.add_page_break()

    # ---- 14 Governance
    h(doc, "14. Cross-center governance, DBI scoring hooks, and references", level=1)
    h(doc, "14.1 Governance", level=2)
    bullets(
        doc,
        [
            "Single study owner maintains mapping table: local protocol name → semantic role → protocol version.",
            "Quarterly audit: sample N exams per site, run automated DBI validator, review failures with site MR lead.",
            "Change requests require signed amendment and version bump.",
        ],
    )
    h(doc, "14.2 DBI component mapping (illustrative)", level=2)
    checklist_table(
        doc,
        ["DBI component", "Modality-specific emphasis"],
        [
            ("Metadata completeness", "Section checklists above"),
            ("Protocol conformity", "protocol_version token + parameter tolerances"),
            ("Gradient integrity (DWI)", "b-table and volume counts"),
            ("Spatial consistency", "spacing, FOV, unexpected reformat"),
            ("Naming compliance", "role tokens and derivative suffixes"),
            ("Drift control", "phantom QC + version discipline"),
        ],
    )
    h(doc, "14.3 External references (informative)", level=2)
    bullets(
        doc,
        [
            "DICOM Standard — https://www.dicomstandard.org/current/",
            "BIDS specification (MRI) — https://bids-specification.readthedocs.io/",
            "BIDS extension proposals (e.g., advanced DWI BEP033) — https://bids.neuroimaging.io/extensions/",
            "FAIR principles — https://www.go-fair.org/fair-principles/",
            "NIH Imaging Data Commons — https://imaging.datacommons.cancer.gov/",
        ],
    )
    p(
        doc,
        "End of document.",
    )

    out = (
        "/mnt/nfs/home/urmc-sh.rochester.edu/pndagiji/Documents/vecta/"
        "AI_Ready_Acquisition_Standards_by_Modality.docx"
    )
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
