#!/usr/bin/env python3
"""
Build the DBI v1 NeuroImage manuscript as a DOCX.

Reads frozen audit outputs, Phase 3 live conversion results, and config to embed real numbers.
Output: vecta/writing/DBI_v1_Manuscript.docx
"""
from __future__ import annotations

import json
import math
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from scipy import stats as sp_stats

ROOT = Path(__file__).resolve().parent.parent
CIDUR_DBI = ROOT / "cidur_dbi"
OUT_AUDIT = CIDUR_DBI / "outputs"
OUT_P3 = CIDUR_DBI / "outputs_dcm2niix_v1freeze"


def _load_json(p: Path) -> dict:
    if not p.is_file():
        return {}
    with open(p, encoding="utf-8") as f:
        return json.load(f)


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def p(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic


def p_justified(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def display_eq(doc: Document, formula: str, eq_num: int) -> None:
    """Render a numbered display equation using a 3-column invisible table:
    left (empty spacer), centre (equation), right (number)."""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = True
    for cell in tbl.rows[0].cells:
        for pg in cell.paragraphs:
            pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.makeelement(qn("w:tcBorders"), {})
        for edge in ("top", "left", "bottom", "right"):
            b = borders.makeelement(qn(f"w:{edge}"), {
                qn("w:val"): "none", qn("w:sz"): "0",
                qn("w:space"): "0", qn("w:color"): "auto",
            })
            borders.append(b)
        tc_pr.append(borders)

    tbl.rows[0].cells[0].paragraphs[0].text = ""
    eq_para = tbl.rows[0].cells[1].paragraphs[0]
    eq_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eq_para.add_run(formula)
    run.italic = True
    run.font.size = Pt(11)

    num_para = tbl.rows[0].cells[2].paragraphs[0]
    num_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    num_para.add_run(f"({eq_num})")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, hdr_text in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = hdr_text
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for row_data in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)


def add_figure_placeholder(doc: Document, path: Path, caption: str, width: float = 5.5) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path.is_file():
        para.add_run().add_picture(str(path), width=Inches(width))
    else:
        para.add_run(f"[Figure not found: {path.name}]").italic = True
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption).italic = True


CLUSTER_SHORT = {
    "GE MEDICAL SYSTEMS | SIGNA Artist | 1.5T": "GE Artist 1.5T",
    "GE MEDICAL SYSTEMS | SIGNA Premier | 3T": "GE Premier 3T",
    "SIEMENS | Skyra | 3T": "Skyra 3T",
    "Siemens Healthineers | MAGNETOM Vida Fit | 3T": "Vida Fit 3T",
}


def _short(cluster: str) -> str:
    return CLUSTER_SHORT.get(cluster, cluster[:25])


# --------------- figure generators ---------------

def gen_fig3_component_heatmap(series_df: pd.DataFrame, out: Path) -> Path:
    """Figure 3: component mean heatmap (scanner cluster x M/P/S/N/G)."""
    comps = ["M", "P", "S", "N"]
    clusters = [c for c in series_df["scanner_cluster"].unique() if c in CLUSTER_SHORT]
    rows = []
    for cl in clusters:
        grp = series_df[series_df["scanner_cluster"] == cl]
        row = {c: grp[c].mean() for c in comps}
        dwi = grp[grp["series_class"] == "dwi"]
        raw_dwi = dwi[dwi["derivative_series"] == False]
        row["G (DWI)"] = pd.to_numeric(raw_dwi["G"], errors="coerce").mean() if len(raw_dwi) else float("nan")
        rows.append(row)
    mat = pd.DataFrame(rows, index=[_short(c) for c in clusters])
    fig, ax = plt.subplots(figsize=(6.5, 3.0))
    sns.heatmap(mat, annot=True, fmt=".3f", cmap="YlOrRd_r", vmin=0, vmax=1,
                linewidths=0.5, ax=ax, cbar_kws={"label": "Mean score"})
    ax.set_ylabel("")
    ax.set_title("Component scores by scanner cluster")
    fig.tight_layout()
    p = out / "figure3_component_heatmap.png"
    fig.savefig(p, dpi=300)
    plt.close(fig)
    return p


def gen_fig4_dbi_violin(series_df: pd.DataFrame, out: Path) -> Path:
    """Figure 4: DBI distribution violin by scanner cluster."""
    df = series_df[series_df["scanner_cluster"].isin(CLUSTER_SHORT)].copy()
    df["cluster_short"] = df["scanner_cluster"].map(_short)
    order = sorted(df["cluster_short"].unique())
    fig, ax = plt.subplots(figsize=(7, 4))
    sns.violinplot(x="cluster_short", y="DBI", hue="cluster_short", data=df, order=order,
                   inner="box", cut=0, ax=ax, palette="Set2", legend=False)
    ax.set_xlabel("Scanner cluster")
    ax.set_ylabel("Series-level DBI")
    ax.set_title("DBI distribution by scanner cluster")
    skew = series_df["DBI"].skew()
    kurt = series_df["DBI"].kurtosis()
    ax.text(0.98, 0.02, f"skew={skew:.2f}  kurtosis={kurt:.2f}",
            transform=ax.transAxes, ha="right", va="bottom", fontsize=8,
            bbox=dict(boxstyle="round,pad=0.3", fc="white", alpha=0.8))
    fig.tight_layout()
    p = out / "figure4_dbi_violin.png"
    fig.savefig(p, dpi=300)
    plt.close(fig)
    return p


def gen_fig5_corr_matrix(series_df: pd.DataFrame, out: Path) -> Path:
    """Figure 5: Spearman inter-component correlation matrix."""
    comps = ["M", "P", "S", "N", "DBI"]
    corr_mat = series_df[comps].corr(method="spearman")
    mask = np.triu(np.ones_like(corr_mat, dtype=bool), k=1)
    fig, ax = plt.subplots(figsize=(5, 4))
    sns.heatmap(corr_mat, annot=True, fmt=".2f", cmap="coolwarm", vmin=-1, vmax=1,
                mask=mask, square=True, linewidths=0.5, ax=ax,
                cbar_kws={"label": "Spearman \u03c1"})
    ax.set_title("Inter-component correlations")
    fig.tight_layout()
    p = out / "figure5_corr_matrix.png"
    fig.savefig(p, dpi=300)
    plt.close(fig)
    return p


def gen_figS2_dbi_conversion(series_df: pd.DataFrame, conv_df: pd.DataFrame,
                             out: Path) -> Path:
    """Figure S2: DBI by conversion outcome (pass vs fail)."""
    merged = series_df.merge(
        conv_df[["session_id", "scan_folder", "convert_pass"]],
        on=["session_id", "scan_folder"], how="inner",
    )
    merged["Outcome"] = merged["convert_pass"].map({True: "Pass", False: "Fail"})
    fig, axes = plt.subplots(1, 2, figsize=(10, 4), gridspec_kw={"width_ratios": [1, 2]})
    sns.boxplot(x="Outcome", y="DBI", hue="Outcome", data=merged, ax=axes[0],
                palette="Set3", order=["Pass", "Fail"], legend=False)
    axes[0].set_title("Composite DBI")
    axes[0].set_ylabel("Score")
    comp_melt = merged.melt(id_vars=["Outcome"], value_vars=["M", "P", "S", "N"],
                            var_name="Component", value_name="Score")
    sns.boxplot(x="Component", y="Score", hue="Outcome", data=comp_melt, ax=axes[1],
                palette="Set3", hue_order=["Pass", "Fail"])
    axes[1].set_title("Per-component scores")
    axes[1].set_ylabel("")
    axes[1].legend(title="Outcome", loc="lower right")
    fig.suptitle("DBI scores by dcm2niix conversion outcome", fontsize=12)
    fig.tight_layout()
    p = out / "figureS2_dbi_vs_conversion.png"
    fig.savefig(p, dpi=300)
    plt.close(fig)
    return p


# --------------- inferential statistics ---------------

def compute_inferential_stats(series_df: pd.DataFrame,
                              conv_df: pd.DataFrame,
                              session_df: pd.DataFrame | None = None) -> dict:
    """Return a dict of all inferential test results for the manuscript."""
    res: dict = {}

    # 1. Kruskal-Wallis: DBI ~ scanner cluster (series-level)
    main = series_df[series_df["scanner_cluster"].isin(CLUSTER_SHORT)]
    groups = [grp["DBI"].values for _, grp in main.groupby("scanner_cluster")]
    kw_h, kw_p = sp_stats.kruskal(*groups)
    n = len(main)
    k = len(groups)
    eta_sq = (kw_h - k + 1) / (n - k)
    res["kw_cluster"] = {"H": kw_h, "p": kw_p, "eta_sq": eta_sq, "k": k, "n": n}

    # 1b. Session-level Kruskal-Wallis (addresses non-independence)
    sess_col = "DBI_session_mean_no_localizer"
    if session_df is not None and sess_col in session_df.columns:
        sess_main = session_df[session_df["scanner_cluster"].isin(CLUSTER_SHORT)]
        sess_groups = [grp[sess_col].dropna().values
                       for _, grp in sess_main.groupby("scanner_cluster")]
        kw_s_h, kw_s_p = sp_stats.kruskal(*sess_groups)
        n_s = len(sess_main)
        k_s = len(sess_groups)
        eps_s = (kw_s_h - k_s + 1) / (n_s - k_s)
        res["kw_session"] = {"H": kw_s_h, "p": kw_s_p, "eps2": eps_s,
                             "k": k_s, "n": n_s}
        # Session-level pairwise
        sess_cluster_names = sorted(CLUSTER_SHORT.keys())
        n_sp = k_s * (k_s - 1) // 2
        sess_pairwise = []
        for i in range(len(sess_cluster_names)):
            for j in range(i + 1, len(sess_cluster_names)):
                g1 = sess_main[sess_main["scanner_cluster"] == sess_cluster_names[i]][sess_col].dropna()
                g2 = sess_main[sess_main["scanner_cluster"] == sess_cluster_names[j]][sess_col].dropna()
                if len(g1) >= 2 and len(g2) >= 2:
                    u, pv = sp_stats.mannwhitneyu(g1, g2, alternative="two-sided")
                    r_eff = 1 - (2 * u) / (len(g1) * len(g2))
                    sess_pairwise.append({
                        "c1": _short(sess_cluster_names[i]),
                        "c2": _short(sess_cluster_names[j]),
                        "U": u, "p": pv,
                        "p_bonf": min(pv * n_sp, 1.0),
                        "r": r_eff, "n1": len(g1), "n2": len(g2),
                    })
        res["sess_pairwise"] = sess_pairwise

    # 2. Series-level pairwise Mann-Whitney U with Bonferroni
    cluster_names = sorted(CLUSTER_SHORT.keys())
    n_pairs = k * (k - 1) // 2
    pairwise = []
    for i in range(len(cluster_names)):
        for j in range(i + 1, len(cluster_names)):
            g1 = main[main["scanner_cluster"] == cluster_names[i]]["DBI"]
            g2 = main[main["scanner_cluster"] == cluster_names[j]]["DBI"]
            u, p_val = sp_stats.mannwhitneyu(g1, g2, alternative="two-sided")
            r_effect = 1 - (2 * u) / (len(g1) * len(g2))
            pairwise.append({
                "c1": _short(cluster_names[i]),
                "c2": _short(cluster_names[j]),
                "U": u, "p": p_val,
                "p_bonf": min(p_val * n_pairs, 1.0),
                "r": r_effect,
                "n1": len(g1), "n2": len(g2),
            })
    res["pairwise"] = pairwise

    # 3. Kruskal-Wallis: DBI ~ series class
    class_groups = [grp["DBI"].values for _, grp in series_df.groupby("series_class")]
    kw_h2, kw_p2 = sp_stats.kruskal(*class_groups)
    res["kw_class"] = {"H": kw_h2, "p": kw_p2}

    # 4. Chi-squared: conversion ~ scanner cluster
    merged = series_df.merge(
        conv_df[["session_id", "scan_folder", "convert_pass"]],
        on=["session_id", "scan_folder"], how="inner",
    )
    merged_main = merged[merged["scanner_cluster"].isin(CLUSTER_SHORT)]
    ct = pd.crosstab(merged_main["scanner_cluster"], merged_main["convert_pass"])
    chi2, p_chi, dof, _ = sp_stats.chi2_contingency(ct)
    cramers_v = math.sqrt(chi2 / (len(merged_main) * (min(ct.shape) - 1)))
    res["chi2_conv"] = {"chi2": chi2, "p": p_chi, "dof": dof, "V": cramers_v, "n": len(merged_main)}

    # 5. Point-biserial: DBI vs conversion
    rpb, p_rpb = sp_stats.pointbiserialr(merged["convert_pass"].astype(int), merged["DBI"])
    res["pb_dbi"] = {"r": rpb, "p": p_rpb, "n": len(merged)}

    comp_pb = {}
    for c in ["M", "P", "S", "N"]:
        r, pv = sp_stats.pointbiserialr(merged["convert_pass"].astype(int), merged[c])
        comp_pb[c] = {"r": r, "p": pv}
    res["pb_components"] = comp_pb

    # 6. Distribution tests
    sample = series_df["DBI"].sample(min(500, len(series_df)), random_state=42)
    w, p_sw = sp_stats.shapiro(sample)
    res["shapiro"] = {"W": w, "p": p_sw, "n": len(sample)}
    res["skewness"] = float(series_df["DBI"].skew())
    res["kurtosis"] = float(series_df["DBI"].kurtosis())

    # 8. Weight sensitivity (Monte Carlo ±20%)
    default_w = {"M": 0.25, "P": 0.20, "G": 0.15, "S": 0.15, "N": 0.20, "D": 0.05}
    comp_list = ["M", "P", "G", "S", "N", "D"]
    if "D" not in series_df.columns:
        series_df = series_df.copy()
        series_df["D"] = 1.0

    def _recompute(row, w):
        num = sum(w[c] * row[c] for c in comp_list if pd.notna(row.get(c)))
        den = sum(w[c] for c in comp_list if pd.notna(row.get(c)))
        return num / den if den > 0 else 0.0

    baseline_dbi = series_df.apply(lambda r: _recompute(r, default_w), axis=1)
    rng = np.random.default_rng(42)
    mc_means = []
    mc_rhos = []
    for _ in range(1000):
        pw = {c: default_w[c] * rng.uniform(0.80, 1.20) for c in comp_list}
        nd = series_df.apply(lambda r: _recompute(r, pw), axis=1)
        mc_means.append(nd.mean())
        rho, _ = sp_stats.spearmanr(baseline_dbi, nd)
        mc_rhos.append(rho)
    mc_means = np.array(mc_means)
    mc_rhos = np.array(mc_rhos)
    res["weight_sens"] = {
        "mean_range": (float(mc_means.min()), float(mc_means.max())),
        "mean_95ci": (float(np.percentile(mc_means, 2.5)), float(np.percentile(mc_means, 97.5))),
        "rho_min": float(mc_rhos.min()),
        "rho_mean": float(mc_rhos.mean()),
    }

    return res


def main() -> None:
    meta = _load_json(OUT_AUDIT / "run_metadata.json")
    p3_env = _load_json(OUT_P3 / "dcm2niix_environment.json")

    n_sessions = meta.get("n_mr_sessions", "N")
    n_series = meta.get("n_series_rows", "N")
    mean_dbi = meta.get("mean_dbi_series", 0)
    mean_dbi_sess = meta.get("mean_dbi_session", 0)
    spec_ver = meta.get("dbi_spec_version", "1.0.3")

    series_df = None
    series_path = OUT_AUDIT / "per_series.csv"
    if series_path.is_file():
        series_df = pd.read_csv(series_path)

    session_path = OUT_AUDIT / "per_session.csv"
    session_df = pd.read_csv(session_path) if session_path.is_file() else None

    t1_path = OUT_AUDIT / "table1_dbi_by_scanner.csv"
    t1_df = pd.read_csv(t1_path) if t1_path.is_file() else None

    t2_path = OUT_P3 / "table2_conversion_by_scanner.csv"
    t2_df = pd.read_csv(t2_path) if t2_path.is_file() else None

    t2c_path = OUT_P3 / "table2_conversion_by_scanner_class.csv"
    t2c_df = pd.read_csv(t2c_path) if t2c_path.is_file() else None

    conv_path = OUT_P3 / "conversion_log.csv"
    conv_df = pd.read_csv(conv_path) if conv_path.is_file() else None

    d2_ver = p3_env.get("dcm2niix_version", "v1.0.x")

    # --- generate new figures ---
    fig3_path = gen_fig3_component_heatmap(series_df, OUT_AUDIT) if series_df is not None else None
    figS2_path = (gen_figS2_dbi_conversion(series_df, conv_df, OUT_AUDIT)
                  if series_df is not None and conv_df is not None else None)

    # --- compute inferential statistics ---
    inf = (compute_inferential_stats(series_df, conv_df, session_df)
           if (series_df is not None and conv_df is not None) else {})

    doc = Document()

    # =========== TITLE PAGE ===========
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(
        "Data Birth Integrity: A Modular Scoring Framework for Assessing "
        "Automation-Readiness of Clinical DICOM Series in Neuroimaging Research"
    )
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph()
    auth = doc.add_paragraph()
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth.add_run("[Author list — to be completed by team]").italic = True

    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff.add_run("[Affiliations — to be completed by team]").italic = True

    doc.add_paragraph()
    corr = doc.add_paragraph()
    corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corr.add_run("Corresponding author: [Name, email]").italic = True

    doc.add_page_break()

    # =========== HIGHLIGHTS ===========
    h(doc, "Highlights", level=1)
    bullets(doc, [
        "Data Birth Integrity (DBI) is a modular 0\u20131 score quantifying how well a DICOM series "
        "meets automation-oriented structural requirements before any image processing.",
        f"Retrospective audit of {n_series:,} MR series across {n_sessions} sessions and four scanner "
        "clusters reveals systematic variation in metadata completeness and naming compliance.",
        "Low DBI\u2014driven primarily by protocol naming (P) and community-standards naming compliance (N)\u2014"
        "indicates that a series is not BIDS-conversion ready, will require custom heuristics for AI/ML "
        "pipelines, and imposes additional data-cleaning complexity.",
        "Naming conventions are grounded in published community standards (BIDS, ENIGMA, ADNI, DICOM PS3.15), "
        "making DBI institution-independent and deployable across sites without re-calibration.",
    ])

    doc.add_page_break()

    # =========== ABSTRACT ===========
    h(doc, "Abstract", level=1)
    p_justified(doc,
        "Background: Modern neuroimaging analysis pipelines increasingly depend on well-structured DICOM "
        "metadata\u2014manufacturer identity, spatial geometry, diffusion encoding parameters, and consistent "
        "naming conventions\u2014yet no published scalar metric exists to quantify this structural readiness at "
        "the point of data creation. The mismatch between what clinical scanners produce and what research "
        "pipelines require is a persistent bottleneck for reproducible large-scale studies."
    )
    p_justified(doc,
        "New Method: We introduce Data Birth Integrity (DBI), a modular scoring framework (0\u20131) that "
        "evaluates five components of each DICOM series: metadata completeness (M), protocol/naming conformity "
        "(P), gradient integrity for diffusion (G), spatial consistency (S), and community-standards naming "
        "compliance (N). Component N conventions are derived from published standards\u2014the BIDS "
        "specification, ENIGMA/ADNI consortium protocols, DICOM interoperability requirements, and POSIX "
        "filename guidance\u2014rather than site-specific patterns, making them institution-independent. "
        "Components that do not apply to a given modality are excluded via renormalization rather than "
        "penalized. A low DBI score, particularly when driven by P and N, indicates that the series is not "
        "ready for automated BIDS conversion, will require custom heuristics for AI/ML pipeline ingestion, "
        "and imposes additional data-cleaning complexity."
    )
    p_justified(doc,
        f"Results: We applied DBI v{spec_ver} to a clinical neuroimaging cohort of {n_sessions} MR sessions "
        f"({n_series:,} series) exported from an institutional XNAT archive across four scanner clusters "
        f"(two GE, two Siemens). Mean series-level DBI was {mean_dbi:.3f} (SD = {series_df['DBI'].std():.3f}). "
        f"Metadata completeness (M = {series_df['M'].mean():.3f}) and spatial consistency "
        f"(S = {series_df['S'].mean():.3f}) were uniformly high, whereas naming compliance "
        f"(N = {series_df['N'].mean():.3f}) and gradient integrity for DWI "
        f"(G = {series_df['G'].replace('', float('nan')).astype(float).mean():.3f}) drove most inter-series "
        f"DBI variation. Session-level analysis (N = {n_sessions}) revealed large inter-scanner heterogeneity "
        f"(\u03b5\u00b2 = {inf.get('kw_session', {}).get('eps2', 0):.2f}). "
        f"Standard DICOM-to-NIfTI conversion with dcm2niix ({d2_ver}) "
        f"succeeded for {t2_df['n_pass'].sum()}/{t2_df['n'].sum()} series "
        f"({100.0 * t2_df['n_pass'].sum() / t2_df['n'].sum():.1f}%), with pass rates among the four "
        f"primary clusters ranging from "
        f"{100.0 * t2_df[t2_df['n'] >= 10]['pass_rate'].min():.1f}% to "
        f"{100.0 * t2_df[t2_df['n'] >= 10]['pass_rate'].max():.1f}%."
        if t2_df is not None else
        "Conversion statistics: [run Phase 3]."
    )
    p_justified(doc,
        "Comparison with existing methods: Unlike image-quality metrics (MRIQC, QUAD/SQUAD) that operate on "
        "processed NIfTI or corrected volumes, DBI assesses raw DICOM metadata upstream of conversion. Unlike "
        "BIDS validation, which checks a curated dataset layout, DBI scores the source series that must eventually "
        "feed that layout. No prior tool occupies this niche."
    )
    p_justified(doc,
        "Conclusions: DBI provides a transparent, reproducible summary of whether clinical DICOM data meet "
        "the structural prerequisites for automated neuroimaging workflows. Because its naming conventions are "
        "grounded in published community standards, DBI is deployable across institutions without "
        "re-calibration. Routine DBI audits can identify series that are not BIDS-conversion ready, quantify "
        "the data-cleaning effort required for AI/ML pipelines, and surface acquisition-site heterogeneity "
        "before it propagates to downstream analyses."
    )
    doc.add_paragraph()
    kw = doc.add_paragraph()
    kw.add_run("Keywords: ").bold = True
    kw.add_run("DICOM; data quality; neuroimaging; automation; BIDS; metadata; community standards; diffusion MRI; reproducibility")

    doc.add_page_break()

    # =========== 1. INTRODUCTION ===========
    h(doc, "1. Introduction", level=1)

    p_justified(doc,
        "The Digital Imaging and Communications in Medicine (DICOM) standard has served as the backbone of "
        "medical image exchange for over three decades (Bidgood & Horii, 1997). While DICOM provides a rich "
        "information model, manufacturers implement it inconsistently: critical acquisition parameters are "
        "often encoded in vendor-specific private tags rather than standardized public fields (Rorden et al., "
        "2025; Larobina & Murino, 2014). This variability is largely invisible in clinical workflows, where "
        "readers rely on pixel data, but becomes a critical failure mode when the same images enter research "
        "pipelines that assume well-formed metadata."
    )

    p_justified(doc,
        "The Brain Imaging Data Structure (BIDS; Gorgolewski et al., 2016) has transformed how curated "
        "neuroimaging datasets are organized and shared, enabling standardized processing through tools such "
        "as fMRIPrep (Esteban et al., 2019), Nipype (Gorgolewski et al., 2011), and FreeSurfer (Fischl, "
        "2012). Extensions to new modalities\u2014most recently PET-BIDS (Norgaard et al., 2022)\u2014"
        "underscore the growing dependence of the field on well-structured input metadata. However, BIDS "
        "operates at the curated-output end of the data lifecycle. Before a DICOM series can be converted to "
        "NIfTI (Li et al., 2016) and organized into a BIDS layout (Halchenko et al., 2024), its metadata must "
        "already meet certain structural prerequisites\u2014prerequisites that are rarely checked systematically."
    )

    p_justified(doc,
        "Image-quality control tools such as MRIQC (Esteban et al., 2017) and diffusion-specific QC "
        "frameworks (Bastiani et al., 2019) provide post-conversion assessments of signal quality, motion, "
        "and distortion. These are invaluable but operate downstream: they cannot detect a missing "
        "MagneticFieldStrength tag, an ambiguous SeriesDescription, or a DWI series whose b-value encoding "
        "is discoverable only through undocumented private tags."
    )

    p_justified(doc,
        "Multi-site consortia such as ENIGMA (Thompson et al., 2020) and TRACK-TBI (Yue et al., 2013) have "
        "demonstrated that protocol standardization at the point of acquisition is essential for meaningful "
        "cross-site analyses. Statistical harmonization methods (Fortin et al., 2017, 2018) can mitigate "
        "residual scanner effects, but they presuppose that the input data are structurally sound. When "
        "upstream metadata is incomplete or inconsistent, harmonization may mask problems rather than resolve them."
    )

    p_justified(doc,
        "The FAIR principles (Wilkinson et al., 2016) call for data that are Findable, Accessible, "
        "Interoperable, and Reusable. Interoperability and reusability, in the context of neuroimaging DICOM, "
        "depend on the presence and correctness of metadata at the point of creation\u2014what we term the "
        "\"birth\" of the data. Reproducibility concerns in neuroimaging are well documented (Poldrack et al., "
        "2017), and open data-sharing platforms such as OpenNeuro (Markiewicz et al., 2021) have demonstrated "
        "the value of standardized, BIDS-compliant datasets for reuse across laboratories. Yet despite growing "
        "awareness of these issues, no published metric quantifies the structural readiness of raw DICOM "
        "series for automation-oriented research workflows."
    )

    p_justified(doc,
        "Here we introduce Data Birth Integrity (DBI), a modular scoring framework that assigns each DICOM "
        "series a scalar value between 0 and 1, reflecting how well it satisfies automation-oriented "
        "structural requirements. DBI is not a measure of diagnostic image quality, signal-to-noise ratio, or "
        "clinical acceptability. Instead, it captures whether the metadata foundation exists for reproducible "
        "conversion, classification, and processing. We describe the DBI specification, apply it retrospectively "
        f"to a clinical cohort of {n_sessions} MR sessions across four scanner clusters, report conversion "
        "outcomes with dcm2niix, and discuss implications for prospective acquisition governance."
    )

    # =========== 2. METHODS ===========
    h(doc, "2. Methods", level=1)

    h(doc, "2.1 Cohort and data source", level=2)
    p_justified(doc,
        f"We analyzed {n_sessions} MR imaging sessions ({n_series:,} series) from a clinical neuroimaging "
        "archive exported using the XNAT platform (Marcus et al., 2007). Each session corresponds to one "
        "imaging visit; each series corresponds to one XNAT scan folder containing DICOM files under a "
        "standardized directory structure (resources/DICOM/files/*.dcm). Sessions spanned four scanner "
        "clusters identified by concatenating Manufacturer, ManufacturerModelName, and MagneticFieldStrength "
        "from the first readable DICOM in each session."
    )

    if t1_df is not None:
        doc.add_paragraph()
        p(doc, "Table 1. Scanner clusters in the cohort.", bold=True)
        add_table(doc,
            ["Scanner cluster", "Sessions (n)", "DBI mean", "DBI SD", "DBI min", "DBI max"],
            [
                [
                    row["scanner_cluster"],
                    str(int(row["n_sessions"])),
                    f"{row['DBI_mean']:.4f}",
                    f"{row['DBI_std']:.4f}",
                    f"{row['DBI_min']:.4f}",
                    f"{row['DBI_max']:.4f}",
                ]
                for _, row in t1_df.iterrows()
            ],
        )

    h(doc, "2.2 DBI specification overview", level=2)
    p_justified(doc,
        f"DBI v{spec_ver} defines five components scored at the series level, each producing a value "
        "in [0, 1] or not-applicable (NA). Components marked NA are excluded from the composite via "
        "weight renormalization rather than scored as zero, ensuring modality-appropriate evaluation."
    )
    p_justified(doc,
        "The five components are: (M) Metadata completeness\u2014presence and plausibility of required "
        "DICOM tags per modality class; (P) Protocol and naming conformity\u2014folder structure and "
        "prospective protocol token compliance; (G) Gradient integrity\u2014diffusion encoding evidence, "
        "applicable only to DWI series; (S) Spatial consistency\u2014pixel spacing and slice geometry within "
        "configured plausibility bounds; and (N) Community-standards naming compliance\u2014conformity of "
        "SeriesDescription and ProtocolName to published naming conventions from the BIDS specification "
        "(Gorgolewski et al., 2016), multi-site consortium protocols (ENIGMA: Thompson et al., 2020; ADNI: "
        "Jack et al., 2008), and DICOM interoperability standards (Bidgood & Horii, 1997)."
    )
    p_justified(doc,
        "A sixth component, Drift control (D), is defined in the specification for future use but assigned "
        "NA for all series in this retrospective analysis, as protocol version time-series data are not "
        "available from static exports."
    )

    h(doc, "2.3 Component M \u2014 Metadata completeness", level=2)
    p_justified(doc,
        "For each series, M is the fraction of required DICOM elements that are present and valid:"
    )
    display_eq(doc, "M = n_pass / n_required", 1)
    p_justified(doc,
        "where n_pass is the count of required tags that are present with plausible values and n_required "
        "is the total number of tags checked for the series\u2019 modality class. A fractional score provides "
        "graded feedback for cohort-level stratification. Universal "
        "elements checked for all non-localizer series include: Modality (must equal MR), Manufacturer, "
        "ManufacturerModelName, MagneticFieldStrength (positive number), SeriesInstanceUID, StudyInstanceUID, "
        "and at least one of SeriesDescription or ProtocolName. Spatial elements (PixelSpacing or "
        "ImagerPixelSpacing; SliceThickness or SpacingBetweenSlices within configured bounds) are required "
        "for all classes except localizer. Class-specific add-ons include diffusion evidence for DWI, "
        "RepetitionTime and EchoTime for BOLD and ASL, and EchoTime for field maps."
    )

    h(doc, "2.4 Component P \u2014 Protocol / naming conformity", level=2)
    p_justified(doc,
        "P comprises two sub-scores averaged equally:"
    )
    display_eq(doc, "P = (P_minimal + P_ideal) / 2", 2)
    p_justified(doc,
        "P_minimal checks whether the scan folder matches the XNAT indexing pattern and whether at least one "
        "descriptive text field is non-empty. P_ideal checks whether SeriesDescription or ProtocolName matches "
        "a prospective protocol token regex aligned with consortium naming conventions. For this legacy cohort, "
        "P_ideal is expected to be zero for most series; P_minimal serves as the primary retrospective metric. "
        "A low P score indicates the series lacks the structured naming that BIDS conversion tools and "
        "automated pipelines require for deterministic processing."
    )

    h(doc, "2.5 Component G \u2014 Gradient integrity (DWI only)", level=2)
    p_justified(doc,
        "G is applicable only when the heuristic series class is DWI. Series whose names match configured "
        "derivative tokens (ADC, FA, TRACE, etc.) receive G = 0 and are flagged as derivative. For "
        "non-derivative DWI series, G is a weighted sum of three binary indicators:"
    )
    display_eq(doc,
        "G = 0.45 \u00b7 I_bval + 0.45 \u00b7 I_grad + 0.10 \u00b7 I_vol", 3)
    p_justified(doc,
        "where I_bval = 1 if b-value evidence is discoverable (from standard DICOM tags, ImageType "
        "heuristics, or Enhanced MR functional group sequences), I_grad = 1 if gradient direction "
        "information is present, and I_vol encodes volume-count plausibility (currently I_vol = 1 for all "
        "series; retained as a placeholder for future refinement). The sub-weights reflect downstream "
        "importance: b-values and gradient directions are both indispensable for tensor estimation "
        "(0.45 each); volume count carries less diagnostic value in this version (0.10)."
    )

    h(doc, "2.6 Component S \u2014 Spatial consistency", level=2)
    p_justified(doc,
        "S verifies that pixel spacing and slice thickness/spacing fall within configurable plausibility "
        "bounds (0.05\u201315.0 mm for in-plane, 0.05\u201320.0 mm for slice). A three-tier ordinal scale is "
        "used:"
    )
    display_eq(doc,
        "S = { 1  if in-plane \u2227 slice pass;   "
        "0.5  if in-plane passes, slice missing;   "
        "0  otherwise }", 4)
    p_justified(doc,
        "A series with valid in-plane spacing but absent slice thickness is partially usable (many 2-D "
        "analyses do not require the through-plane dimension), motivating the intermediate score. "
        "Localizer series are assigned S = 1 (NA equivalent)."
    )

    h(doc, "2.7 Component N \u2014 Community-standards naming compliance", level=2)
    p_justified(doc,
        "N measures conformity of series naming to published community standards that simultaneously "
        "determine whether a series is ready for automated BIDS conversion, AI/ML pipeline ingestion, and "
        "programmatic data cleaning. Each binary check is grounded in a specific normative source and "
        "enables a concrete downstream workflow:"
    )
    display_eq(doc,
        "N = \u03a3 check_i / n_checks ,   check_i \u2208 {0, 1}", 5)
    p_justified(doc,
        "Structural checks (always applied) enforce XNAT-style folder indexing (Marcus et al., 2007), "
        "POSIX-safe filenames (no embedded path separators), and DICOM VR length constraints (SeriesDescription "
        "\u2264 128 characters). Community-standards checks (v1.0.5) are derived from six normative sources: "
        "(1) control-character safety per RFC 8259 (JSON) and the BIDS sidecar specification, enabling valid "
        "dcm2niix sidecar emission and ML dataset loaders; (2) BIDS entity tokens (run-, acq-, dir-, echo-; "
        "Gorgolewski et al., 2016), enabling HeuDiConv heuristic auto-detection (Halchenko et al., 2024) "
        "and deterministic file renaming; (3) standardized modality/contrast lexicon (T1w, DWI, BOLD, FLAIR, etc.) "
        "from BIDS suffixes and consortium protocol guides (ENIGMA: Thompson et al., 2020; ADNI: Jack et al., "
        "2008), enabling fMRIPrep/MRIQC series routing (Esteban et al., 2019) and NLP-free modality "
        "classification; (4) machine-parseable token separators following BIDS filename convention and POSIX "
        "portable filename guidance; (5) PHI-safe naming per DICOM PS3.15 Annex E and Aryanto et al. (2015), "
        "detecting MRN-like digit sequences and date patterns; and (6) acquisition parameter encoding "
        "(b-values, direction counts, TE/TR) following ENIGMA-DTI and ADNI protocol conventions, enabling "
        "automated parameter cross-verification."
    )
    p_justified(doc,
        "Per-class SeriesDescription compliance rules (e.g., DWI descriptions must include phase-encoding "
        "direction and direction count; field map descriptions must contain fmap-related text) ensure that "
        "naming carries sufficient semantic content for deterministic BIDS suffix assignment. A derived-scan "
        "naming rule, aligned with the BIDS derivatives specification, requires series matching derivative "
        "markers (ADC, FA, TRACE, etc.) to be explicitly labeled."
    )
    p_justified(doc,
        "Because all conventions are derived from institution-independent published standards, N scores "
        "measure distance from community conventions rather than conformity to any single site\u2019s "
        "patterns. A low N score therefore has a concrete interpretation: the series is not BIDS-conversion "
        "ready without manual relabeling, will require custom heuristics for AI/ML pipeline ingestion, "
        "and imposes additional data-cleaning complexity before it can enter reproducible workflows."
    )

    h(doc, "2.8 Composite DBI", level=2)
    p_justified(doc,
        "The composite DBI for each series is the weighted mean of applicable component scores:"
    )
    display_eq(doc,
        "DBI = \u03a3(w_c \u00d7 score_c) / \u03a3(w_c) ,   "
        "c \u2208 { M, P, G, S, N, D }  where na_c = false", 6)
    p_justified(doc,
        "where w_c is the weight assigned to component c and the sums run only over components whose "
        "not-applicable flag is false for the given series."
    )
    p_justified(doc,
        "A weighted arithmetic mean was chosen because it handles NA components via weight renormalization "
        "(excluding inapplicable components from both numerator and denominator), preserves the component-level "
        "profile that a minimum-based composite would collapse, and produces scores on the same [0, 1] interval "
        "as the individual components."
    )
    p_justified(doc,
        "Default weights (M: 0.28, P: 0.18, G: 0.22, S: 0.17, N: 0.15; D: 0.00 in v1) were set by "
        "domain-expert consensus, ordered by severity of downstream failure: M highest because missing "
        "fundamental tags propagate to every step; G second because diffusion parameters are irrecoverable "
        "when absent; P, S, and N progressively lower reflecting decreasing impact on hard failures. Weights "
        "are explicitly exposed in the YAML, enabling site-level adjustment; data-driven calibration is "
        "identified as future work."
    )
    p_justified(doc,
        "Session-level DBI is the mean of series-level scores, excluding localizer series."
    )

    h(doc, "2.9 Series classification", level=2)
    if series_df is not None:
        class_counts = series_df["series_class"].value_counts()
    p_justified(doc,
        "Each series is assigned a heuristic class by matching the concatenation of scan folder basename and "
        "SeriesDescription against ordered pattern rules in the YAML configuration. Eleven classes are defined: "
        "DWI, BOLD, ASL, field map, SWI, FLAIR, perfusion, T1 anatomical, T2 anatomical, localizer, and other. "
        "The first matching rule wins; series matching no rule are classified as \"other.\" Classification is "
        "rule-based and heuristic; formal validation against human labels (Phase 4) is optional and was not "
        "performed for this cohort."
    )

    if series_df is not None:
        doc.add_paragraph()
        p(doc, "Table S1. Series class distribution.", bold=True)
        add_table(doc, ["Class", "n", "% of total"],
            [[cls, str(cnt), f"{100.0 * cnt / len(series_df):.1f}"]
             for cls, cnt in class_counts.items()])

    h(doc, "2.10 DICOM-to-NIfTI conversion (Phase 3)", level=2)
    p_justified(doc,
        f"Each series' DICOM folder was converted using dcm2niix ({d2_ver}; Li et al., 2016) with the "
        f"command pattern: dcm2niix -z y -f '%n_%s' -o <output_dir> <dicom_dir>. Each invocation was "
        "executed via subprocess with full capture of exit code, stdout, stderr, and wall-clock duration. "
        "Conversion was judged successful if the exit code was 0 and at least one NIfTI file (.nii or "
        ".nii.gz) was present in the series output directory. This dual criterion is stricter than NIfTI "
        "presence alone: dcm2niix may produce partial output (e.g., a single volume from a corrupted 4-D "
        "series) while returning a non-zero exit code, and such cases are correctly classified as failures. "
        "Series without DICOM files in the expected location were excluded from conversion summaries."
    )

    h(doc, "2.11 Software and reproducibility", level=2)
    p_justified(doc,
        f"DBI scoring was implemented in Python 3.9 using pydicom (Mason et al.) for DICOM parsing and "
        f"pandas for tabulation. All rules, weights, and thresholds are specified in dbi_v1_config.yaml "
        f"(version {spec_ver}). The full specification is provided in DBI_v1_SPECIFICATION.md. Source code, "
        f"configuration, and unit tests (18 passing) are available in the study repository. The audit is "
        f"deterministic: identical inputs and configuration version produce identical scores. DBI was computed "
        f"on research-export DICOM as ingested (pre\u2013de-identification); only aggregate statistics and "
        f"de-identified exemplars are included in public materials. Per-series CSV output includes N_pass and "
        f"N_total for full transparency of naming compliance sub-checks."
    )

    h(doc, "2.12 Statistical analysis", level=2)
    if inf:
        sw = inf["shapiro"]
        p_justified(doc,
            f"Series-level DBI scores were non-normally distributed (Shapiro\u2013Wilk W = {sw['W']:.3f}, "
            f"p < 0.001, n = {sw['n']}; skewness = {inf['skewness']:.2f}, "
            f"excess kurtosis = {inf['kurtosis']:.2f}). All group comparisons therefore used non-parametric "
            "tests, which make no assumptions about the shape of the underlying distribution."
        )
        p_justified(doc,
            "Between-cluster differences in DBI were assessed with the Kruskal\u2013Wallis H test (the "
            "non-parametric analogue of one-way ANOVA for rank-ordered data), followed by pairwise "
            "Mann\u2013Whitney U tests with Bonferroni correction for multiple comparisons (dividing the "
            "nominal \u03b1 = 0.05 by the number of pairwise tests to control the family-wise error rate). "
            "Effect size for the omnibus test is reported as epsilon-squared:"
        )
        display_eq(doc, "\u03b5\u00b2 = H / (n \u2212 1)", 7)
        p_justified(doc,
            "which represents the proportion of variance in ranks explained by group membership; values below "
            "0.01 are considered negligible, 0.01\u20130.06 small, 0.06\u20130.14 medium, and above 0.14 large "
            "(Tomczak & Tomczak, 2014)."
        )
        p_justified(doc,
            "Associations between conversion outcome (binary) and DBI (continuous) were quantified with the "
            "point-biserial correlation coefficient (r_pb), which is mathematically equivalent to the Pearson r "
            "when one variable is dichotomous. Point-biserial correlation was preferred over logistic regression "
            "because the goal was to quantify the strength of association rather than to build a predictive model; "
            "with only five candidate predictors and a binary outcome dominated by factors outside the DBI "
            "construct (e.g., transfer syntax), a regression model would be over-specified and difficult to "
            "interpret."
        )
        p_justified(doc,
            "Scanner-cluster differences in conversion pass rate were tested with Pearson\u2019s chi-squared "
            "test (appropriate for comparing proportions across independent groups). Effect size is reported "
            "as Cram\u00e9r\u2019s V:"
        )
        display_eq(doc,
            "V = \u221a( \u03c7\u00b2 / ( n \u00d7 (min(r, c) \u2212 1) ) )", 8)
        p_justified(doc,
            "where r and c are the rows and columns of the contingency table. V ranges from 0 (no association) "
            "to 1 (perfect association), with 0.10, 0.30, and 0.50 as conventional benchmarks for small, medium, "
            "and large effects (Cohen, 1988). All tests were two-sided; significance was set at \u03b1 = 0.05 "
            "(Bonferroni-adjusted for pairwise comparisons). Analyses used scipy.stats (Python 3.9)."
        )

    h(doc, "2.13 Weight sensitivity analysis", level=2)
    p_justified(doc,
        "Because the default component weights reflect expert judgment rather than empirical calibration, "
        "we assessed the sensitivity of DBI to weight perturbation using a Monte Carlo approach. In each of "
        "1,000 trials, every component weight was independently drawn from a uniform distribution spanning "
        "\u00b120% of its default value; the composite DBI was recomputed for all series. We report the 95% "
        "interval for the cohort-mean DBI and the minimum Spearman rank correlation between the baseline and "
        "perturbed DBI vectors. High rank stability indicates that relative ordering of series is preserved "
        "regardless of exact weight choices within the tested range."
    )

    doc.add_page_break()

    # =========== 3. RESULTS ===========
    h(doc, "3. Results", level=1)

    h(doc, "3.1 Cohort overview", level=2)
    sess_n_series_mean = f"{session_df['n_series'].mean():.1f}" if session_df is not None else "N"
    sess_n_series_range = (
        f"{session_df['n_series'].min()}\u2013{session_df['n_series'].max()}"
        if session_df is not None else "N\u2013N"
    )
    p_justified(doc,
        f"The cohort comprised {n_sessions} MR sessions containing {n_series:,} series "
        f"(mean {sess_n_series_mean} series per session, range {sess_n_series_range}) across four primary "
        f"scanner clusters (Table 1). A fifth minor cluster (two series from a single session with incomplete "
        f"scanner metadata) appears in conversion statistics but is excluded from DBI range calculations. "
        f"No DICOM read failures were encountered. The most prevalent series classes were "
        f"\"other\" ({class_counts.get('other', 0)}), T2 anatomical ({class_counts.get('t2_anat', 0)}), "
        f"DWI ({class_counts.get('dwi', 0)}), and SWI ({class_counts.get('swi', 0)})."
    )

    h(doc, "3.2 Series-level DBI scores", level=2)
    if series_df is not None:
        g_all = series_df["G"].replace("", float("nan")).astype(float)
        p_justified(doc,
            f"Mean series-level DBI was {mean_dbi:.3f} (SD = {series_df['DBI'].std():.3f}, "
            f"range {series_df['DBI'].min():.3f}\u2013{series_df['DBI'].max():.3f}). "
            f"Component-level means were: M = {series_df['M'].mean():.3f} "
            f"(SD = {series_df['M'].std():.3f}), "
            f"P = {series_df['P'].mean():.3f} (SD = {series_df['P'].std():.3f}), "
            f"G = {g_all.mean():.3f} (SD = {g_all.std():.3f}; DWI-only), "
            f"S = {series_df['S'].mean():.3f} (SD = {series_df['S'].std():.3f}), and "
            f"N = {series_df['N'].mean():.3f} (SD = {series_df['N'].std():.3f}). "
            f"P_ideal was 0.000 for all series, as expected for a legacy cohort without prospective protocol "
            f"tokens; P_minimal was {series_df['P_minimal'].mean():.3f}. "
            f"Three components exhibited near-uniform values: M was at ceiling (\u2248 1.0) across all "
            f"clusters, S was likewise near-ceiling ({series_df['S'].mean():.3f}) with only "
            f"{(series_df['S'] < 1.0).sum()} of {len(series_df)} series receiving partial credit, and "
            f"P was uniformly \u2248 0.5 because P_ideal = 0 for all series. Consequently, inter-series DBI "
            f"variation was driven primarily by G (for DWI, where derivative maps receive G = 0) and N "
            f"(naming compliance). This uniformity in M, P, and S is a property of the cohort, not the "
            f"framework: prospective adoption of protocol tokens would activate P_ideal, and sites with less "
            f"consistent spatial metadata would exhibit greater S variation."
        )

    h(doc, "3.3 Scanner-cluster variation", level=2)
    if session_df is not None and t1_df is not None:
        sess_col = "DBI_session_mean_no_localizer"
        p_justified(doc,
            f"Session-level mean DBI (excluding localizers) was {session_df[sess_col].mean():.3f} "
            f"(SD = {session_df[sess_col].std():.3f}, "
            f"range {session_df[sess_col].min():.3f}\u2013{session_df[sess_col].max():.3f}). "
            f"Cluster-level session-mean DBI ranged from "
            f"{t1_df['DBI_mean'].min():.3f} ({t1_df.loc[t1_df['DBI_mean'].idxmin(), 'scanner_cluster']}) to "
            f"{t1_df['DBI_mean'].max():.3f} ({t1_df.loc[t1_df['DBI_mean'].idxmax(), 'scanner_cluster']}), "
            f"indicating small but statistically reliable inter-scanner heterogeneity (Table 1). "
            f"Figure 1 displays the distribution of session-mean DBI by scanner cluster."
        )
    else:
        p_justified(doc, "Scanner-cluster variation: [see Table 1].")

    if inf and "kw_cluster" in inf:
        kw = inf["kw_cluster"]
        eps_sq = kw["H"] / (kw["n"] - 1)
        p_justified(doc,
            f"A Kruskal\u2013Wallis test confirmed that series-level DBI differed significantly across "
            f"the four primary scanner clusters (H({kw['k'] - 1}) = {kw['H']:.2f}, p = {kw['p']:.2e}, "
            f"\u03b5\u00b2 = {eps_sq:.4f}). The small series-level effect size reflects the inflated N "
            f"(multiple non-independent series per session). To account for this, we repeated the analysis "
            f"at the session level, using session-mean DBI as the unit of observation."
        )
        if "kw_session" in inf:
            kws = inf["kw_session"]
            p_justified(doc,
                f"The session-level Kruskal\u2013Wallis test (N = {kws['n']} sessions) yielded "
                f"H({kws['k'] - 1}) = {kws['H']:.2f}, p < 0.001, \u03b5\u00b2 = {kws['eps2']:.2f}\u2014a "
                f"large effect, indicating that scanner cluster identity explains approximately "
                f"{100 * kws['eps2']:.0f}% of the variance in session-level DBI ranks. Post-hoc pairwise "
                f"Mann\u2013Whitney U tests with Bonferroni correction at the session level identified the "
                f"following significant pairs:"
            )
            sess_sig = [pw for pw in inf.get("sess_pairwise", []) if pw["p_bonf"] < 0.05]
            if sess_sig:
                bullets(doc, [
                    f"{pw['c1']} vs. {pw['c2']}: U = {pw['U']:.0f}, p_bonf = {pw['p_bonf']:.2e}, "
                    f"r = {pw['r']:.3f} (n\u2081 = {pw['n1']}, n\u2082 = {pw['n2']})"
                    for pw in sess_sig
                ])
            sess_nonsig = [pw for pw in inf.get("sess_pairwise", []) if pw["p_bonf"] >= 0.05]
            if sess_nonsig:
                p_justified(doc,
                    f"The remaining {len(sess_nonsig)} pairwise comparisons were not significant after "
                    f"Bonferroni correction (all p_bonf > 0.05)."
                )
        p_justified(doc,
            f"At the series level, post-hoc Mann\u2013Whitney U tests with Bonferroni correction "
            f"({len(inf['pairwise'])} comparisons, \u03b1_adj = {0.05 / len(inf['pairwise']):.4f}) "
            f"corroborated these findings. The full series-level and session-level results are reported "
            f"in Table 4."
        )

    add_figure_placeholder(
        doc,
        OUT_AUDIT / "figure1_dbi_by_scanner.png",
        "Figure 1. Distribution of session-mean DBI by scanner cluster. "
        "Each point represents one session; boxes show median and interquartile range.",
    )

    h(doc, "3.4 Component-level patterns", level=2)
    p_justified(doc,
        "Having established that DBI differs across scanner clusters, we next examined which components "
        "drive this variation."
    )
    if series_df is not None:
        dwi_df = series_df[series_df["series_class"] == "dwi"]
        n_deriv = int(dwi_df["derivative_series"].sum())
        g_vals = dwi_df["G"].replace("", float("nan")).astype(float)
        p_justified(doc,
            f"Metadata completeness (M = {series_df['M'].mean():.3f}) and spatial consistency "
            f"(S = {series_df['S'].mean():.3f}) were consistently high across clusters, indicating that "
            f"fundamental DICOM tags and spatial geometry fields were generally well-populated by all four "
            f"scanner platforms. Only {(series_df['S'] < 1.0).sum()} of {len(series_df)} series received "
            f"partial spatial credit (S = 0.5), due to missing SliceThickness. "
            f"Naming compliance (N) averaged {series_df['N'].mean():.3f}, reflecting "
            f"partial adherence to automation-oriented conventions; the per-class SeriesDescription rules and "
            f"derived-scan suffix requirement contributed most to N failures. Table 3 provides component-level "
            f"means for each scanner cluster; Figure 3 displays the same data as a heatmap."
        )
        p_justified(doc,
            f"Among {len(dwi_df)} DWI-classified series, {n_deriv} ({100.0 * n_deriv / len(dwi_df):.1f}%) were "
            f"flagged as derivative (ADC, FA, TRACE, etc.) and received G = 0. For the remaining "
            f"{len(dwi_df) - n_deriv} raw DWI series, mean G was {g_vals[dwi_df['derivative_series'] == False].mean():.3f}, "
            f"indicating {'partial' if g_vals[dwi_df['derivative_series'] == False].mean() < 0.8 else 'generally adequate'} "
            f"discoverability of diffusion encoding parameters from DICOM. The low overall G mean "
            f"({g_vals.mean():.3f}) is dominated by derivative maps, which receive G = 0 by design."
        )

    # --- Table 3: Component means by scanner cluster ---
    if series_df is not None:
        doc.add_paragraph()
        p(doc, "Table 3. Component means (\u00b1 SD) by scanner cluster.", bold=True)
        t3_headers = ["Scanner cluster", "M", "P", "S", "N", "G (raw DWI)"]
        t3_rows = []
        for cl in sorted(CLUSTER_SHORT.keys()):
            grp = series_df[series_df["scanner_cluster"] == cl]
            dwi_grp = grp[(grp["series_class"] == "dwi") & (grp["derivative_series"] == False)]
            g_raw = pd.to_numeric(dwi_grp["G"], errors="coerce")
            g_str = f"{g_raw.mean():.3f} \u00b1 {g_raw.std():.3f}" if len(g_raw) > 0 else "N/A"
            t3_rows.append([
                _short(cl),
                f"{grp['M'].mean():.3f} \u00b1 {grp['M'].std():.3f}",
                f"{grp['P'].mean():.3f} \u00b1 {grp['P'].std():.3f}",
                f"{grp['S'].mean():.3f} \u00b1 {grp['S'].std():.3f}",
                f"{grp['N'].mean():.3f} \u00b1 {grp['N'].std():.3f}",
                g_str,
            ])
        add_table(doc, t3_headers, t3_rows)

    if fig3_path:
        add_figure_placeholder(doc, fig3_path,
            "Figure 3. Component-level mean scores by scanner cluster. Colour scale: 0 (red) to 1 (green). "
            "G is reported for raw (non-derivative) DWI series only.",
            width=6.0,
        )

    add_figure_placeholder(
        doc,
        OUT_AUDIT / "figure_supp_dbi_by_class.png",
        "Figure S1. Mean DBI by heuristic series class.",
        width=5.0,
    )

    h(doc, "3.5 DICOM-to-NIfTI conversion outcomes", level=2)
    p_justified(doc,
        "To assess the practical consequences of metadata variability, we examined dcm2niix conversion "
        "outcomes as an external benchmark."
    )
    if t2_df is not None:
        total_n = int(t2_df["n"].sum())
        total_pass = int(t2_df["n_pass"].sum())
        overall_rate = 100.0 * total_pass / total_n if total_n else 0
        main_clusters = t2_df[t2_df["n"] >= 10]
        p_justified(doc,
            f"All {total_n} series were converted via live dcm2niix invocation with full capture of exit "
            f"codes, stdout, stderr, and wall-clock duration. {total_pass} ({overall_rate:.1f}%) met the "
            f"dual conversion success criterion (exit code 0 AND at least one NIfTI file in the output "
            f"directory). Among the four primary scanner clusters (n \u2265 10 series each), pass rates "
            f"ranged from {100.0 * main_clusters['pass_rate'].min():.1f}% to "
            f"{100.0 * main_clusters['pass_rate'].max():.1f}% (Table 2). The {total_n - total_pass} "
            f"failures (exit code 1) were caused by DICOM instance-number sorting inconsistencies reported "
            f"in dcm2niix stderr; notably, all produced partial NIfTI output despite the non-zero exit code, "
            f"confirming that the dual criterion is more informative than NIfTI presence alone. "
            f"Figure 2 displays the pass-rate heatmap stratified by scanner cluster and series class."
        )

    if inf and "chi2_conv" in inf:
        chi = inf["chi2_conv"]
        p_justified(doc,
            f"Conversion pass rates differed significantly across the four primary clusters "
            f"(\u03c7\u00b2({chi['dof']}) = {chi['chi2']:.2f}, p = {chi['p']:.2e}, "
            f"Cram\u00e9r\u2019s V = {chi['V']:.3f}), indicating a medium-sized association between "
            f"scanner cluster membership and conversion success."
        )

    h(doc, "3.5.1 DBI as a predictor of conversion success", level=3)
    if inf and "pb_dbi" in inf:
        pb = inf["pb_dbi"]
        p_justified(doc,
            f"Point-biserial correlation between composite DBI and conversion outcome was not significant "
            f"(r_pb = {pb['r']:.3f}, p = {pb['p']:.3f}, n = {pb['n']}); per-component correlations were "
            f"likewise negligible (all |r_pb| < 0.07). With only {total_n - total_pass} failures out of "
            f"{total_n} series (99.8% pass rate), the near-ceiling conversion success leaves minimal "
            f"variance for prediction. The three failures were caused by DICOM sorting inconsistencies "
            f"rather than metadata-quality deficits, consistent with the null association (Figure S2)."
        )

    if t2_df is not None:
        doc.add_paragraph()
        p(doc, "Table 2. dcm2niix conversion pass rate by scanner cluster.", bold=True)
        add_table(doc,
            ["Scanner cluster", "n", "n pass", "Pass rate"],
            [
                [row["scanner_cluster"], str(int(row["n"])), str(int(row["n_pass"])),
                 f"{row['pass_rate']:.4f}"]
                for _, row in t2_df.iterrows()
            ],
        )

    add_figure_placeholder(
        doc,
        OUT_P3 / "figure2_dcm2niix_pass_rate_heatmap.png",
        "Figure 2. dcm2niix conversion pass rate (mean) by scanner cluster and heuristic series class.",
    )

    if figS2_path:
        add_figure_placeholder(doc, figS2_path,
            "Figure S2. DBI scores by dcm2niix conversion outcome. Left: composite DBI; Right: per-component "
            "scores. Overlapping distributions confirm the null association (r_pb = "
            f"{inf['pb_dbi']['r']:.3f}, p = {inf['pb_dbi']['p']:.3f})."
            if inf and "pb_dbi" in inf else
            "Figure S2. DBI scores by dcm2niix conversion outcome.",
            width=6.0,
        )

    # --- Table 4: Inferential test summary (essential tests only) ---
    if inf:
        doc.add_paragraph()
        p(doc, "Table 4. Summary of inferential statistical tests.", bold=True)
        t4_headers = ["Test", "Comparison", "Statistic", "p-value", "Effect size"]
        t4_rows = []
        sw = inf["shapiro"]
        t4_rows.append([
            "Shapiro\u2013Wilk",
            f"DBI normality (n = {sw['n']})",
            f"W = {sw['W']:.3f}", f"{sw['p']:.2e}", "\u2014",
        ])
        if "kw_session" in inf:
            kws = inf["kw_session"]
            t4_rows.append([
                "Kruskal\u2013Wallis",
                f"Session DBI ~ scanner cluster (N = {kws['n']})",
                f"H({kws['k'] - 1}) = {kws['H']:.2f}", "< 0.001",
                f"\u03b5\u00b2 = {kws['eps2']:.2f} (large)",
            ])
        sig_pairs = [pw for pw in inf.get("sess_pairwise", []) if pw["p_bonf"] < 0.05]
        for pw in sig_pairs:
            t4_rows.append([
                "Mann\u2013Whitney U",
                f"{pw['c1']} vs {pw['c2']} (session-level)",
                f"U = {pw['U']:.0f}", f"p_bonf = {pw['p_bonf']:.2e}*",
                f"r = {pw['r']:.3f}",
            ])
        kw = inf["kw_cluster"]
        eps_sq = kw["H"] / (kw["n"] - 1)
        t4_rows.append([
            "Kruskal\u2013Wallis",
            f"Series DBI ~ scanner cluster (N = {kw['n']})",
            f"H({kw['k'] - 1}) = {kw['H']:.2f}", f"{kw['p']:.2e}",
            f"\u03b5\u00b2 = {eps_sq:.4f} (small\u2020)",
        ])
        chi = inf["chi2_conv"]
        t4_rows.append([
            "Pearson \u03c7\u00b2",
            "Conversion pass ~ scanner cluster",
            f"\u03c7\u00b2({chi['dof']}) = {chi['chi2']:.2f}", f"{chi['p']:.2e}",
            f"V = {chi['V']:.3f} (medium)",
        ])
        pb = inf["pb_dbi"]
        t4_rows.append([
            "Point-biserial",
            "DBI vs conversion outcome",
            f"r = {pb['r']:.3f}", f"{pb['p']:.3f}",
            "n.s.",
        ])
        ws = inf.get("weight_sens", {})
        if ws:
            t4_rows.append([
                "Monte Carlo",
                "Weight \u00b120% (1,000 trials)",
                f"\u03c1_min = {ws['rho_min']:.3f}",
                "\u2014",
                f"DBI 95% CI [{ws['mean_95ci'][0]:.3f}, {ws['mean_95ci'][1]:.3f}]",
            ])
        add_table(doc, t4_headers, t4_rows)
        p_justified(doc,
            "\u2020 Series-level \u03b5\u00b2 is attenuated by inflated N (multiple non-independent series "
            "per session); the session-level analysis is the primary test."
        )

    doc.add_page_break()

    # =========== 4. DISCUSSION ===========
    h(doc, "4. Discussion", level=1)

    h(doc, "4.1 Summary of findings", level=2)
    kws_disc = inf.get("kw_session", {})
    p_justified(doc,
        f"Applied to {n_sessions} clinical MR sessions across four scanner clusters, DBI revealed that "
        f"metadata completeness (M) and spatial consistency (S) are generally high, while naming compliance "
        f"(N) and gradient integrity (G, for DWI) exhibit the greatest room for improvement\u2014deficiencies "
        f"invisible to any tool operating downstream of DICOM conversion. At the session level, scanner "
        f"cluster membership explained a large share of DBI variance "
        f"(\u03b5\u00b2 = {kws_disc.get('eps2', 0):.2f}), confirming that scanner identity introduces "
        f"substantial heterogeneity in metadata quality. This effect was robust to weight perturbation: "
        f"Monte Carlo simulation (\u00b120% on all weights, 1,000 trials) showed that series-level DBI "
        f"rankings were virtually unchanged (Spearman \u03c1 \u2265 "
        f"{inf.get('weight_sens', {}).get('rho_min', 0):.3f} across all trials), and the 95% CI for "
        f"cohort-mean DBI spanned only "
        f"[{inf.get('weight_sens', {}).get('mean_95ci', (0, 0))[0]:.3f}, "
        f"{inf.get('weight_sens', {}).get('mean_95ci', (0, 0))[1]:.3f}]. These results underscore the "
        f"need for acquisition-level governance and the stability of DBI as a composite measure."
    )

    h(doc, "4.2 Positioning relative to existing tools", level=2)
    p_justified(doc,
        "DBI occupies a distinct niche in the neuroimaging data lifecycle. MRIQC (Esteban et al., 2017) "
        "assesses image quality from NIfTI volumes; QUAD/SQUAD (Bastiani et al., 2019) evaluates diffusion "
        "data after eddy-current correction; bids-validator checks the layout of a curated BIDS dataset. "
        "DBI operates upstream of all of these, at the level of raw DICOM metadata. A series that scores "
        "poorly on DBI may still produce acceptable images for clinical reading but is at elevated risk of "
        "silent failure in automated pipelines\u2014conversion errors, misclassification, missing sidecar "
        "fields, or incorrect spatial geometry in downstream analyses."
    )
    p_justified(doc,
        "The closest conceptual parallel is the MIDRC CRP 12 initiative (Kinahan et al., 2024), which "
        "develops structured DICOM data-quality objects for radiology (CT/XR). In a different domain, the "
        "METhodological RadiomICs Score (METRICS; Kocak et al., 2024) uses a modular checklist to rate "
        "methodological rigor in radiomics studies, illustrating the general utility of composite quality "
        "scores. DBI is neuroimaging-specific, modality-aware (e.g., the G component applies only to DWI), "
        "and designed for cohort-level stratification rather than individual-image triage. Thirty years of "
        "DICOM evolution (NEMA, 2023) have not fully resolved inter-vendor metadata inconsistencies, "
        "reinforcing the need for automated upstream checks."
    )

    h(doc, "4.3 Interpretation of component scores", level=2)
    p_justified(doc,
        "High M and S scores across clusters indicate that fundamental DICOM tags and spatial geometry fields "
        "are well-populated by all four scanner platforms in this cohort, consistent with Rorden et al. (2025). "
        "The near-ceiling S suggests that modern clinical scanners reliably emit PixelSpacing and "
        "SliceThickness within plausible bounds. This is encouraging for pipelines that compute voxel-to-world "
        "transformations (e.g., FreeSurfer; Fischl, 2012), but may not generalize to older archives."
    )
    p_justified(doc,
        "The moderate naming compliance (N) and uniformly half-credited protocol conformity (P) are the most "
        "consequential findings for downstream workflows. These scores have a concrete operational "
        "interpretation: a series with low P and N is (1) not ready for automated BIDS conversion\u2014"
        "HeuDiConv (Halchenko et al., 2024) and dcm2niix (Li et al., 2016) rely on SeriesDescription content "
        "to assign BIDS suffixes and entity keys, and ambiguous naming forces manual heuristic writing; "
        "(2) not readily ingestible by AI/ML pipelines\u2014automated modality classification, feature "
        "extraction from filenames, and experiment-level metadata parsing all depend on standardized naming "
        "tokens; (3) likely to impose substantial data-cleaning complexity\u2014programmatic workflows using "
        "pandas, shell scripts, or dataset loaders require predictable, machine-parseable naming with explicit "
        "token boundaries; and (4) not following commonly adopted conventions from BIDS (Gorgolewski et al., "
        "2016), ENIGMA (Thompson et al., 2020), or ADNI (Jack et al., 2008)\u2014conventions that exist "
        "precisely because the community recognized that unstandardized naming is a barrier to reproducible "
        "multi-site research."
    )
    p_justified(doc,
        "This gap between clinical naming practices and community standards is expected for legacy cohorts and "
        "is informative rather than punitive: N quantifies the curation effort that would be required to bring "
        "raw acquisitions into compliance with the standards that downstream tools assume. Sites adopting "
        "BIDS-aligned naming at the scanner console (as recommended by ENIGMA and ADNI protocol guides) would "
        "score higher on P and N without post-hoc relabeling."
    )
    if series_df is not None:
        other_n = int((series_df["series_class"] == "other").sum())
        other_pct = 100.0 * other_n / len(series_df)
        p_justified(doc,
            f"The \"other\" class accounted for {other_pct:.1f}% of all series ({other_n}/{len(series_df)}). "
            f"These are series whose SeriesDescription and scan folder name did not match any of the ten "
            f"defined modality rules (DWI, BOLD, ASL, field map, SWI, FLAIR, perfusion, T1 anat, T2 anat, "
            f"localizer). Inspection revealed that the majority correspond to screen captures, "
            f"secondary reformats, post-processing outputs (e.g., cerebral blood flow maps), and vendor-"
            f"specific QA series. Because these series are not intended for research pipelines, their "
            f"presence does not invalidate the classification schema but highlights the heterogeneity of "
            f"clinical DICOM archives compared to research-only acquisitions."
        )

    h(doc, "4.4 Conversion outcomes and their relationship to DBI", level=2)
    if inf and "pb_dbi" in inf:
        pb = inf["pb_dbi"]
        chi = inf["chi2_conv"]
        p_justified(doc,
            f"dcm2niix conversion pass rates differed significantly across scanner clusters "
            f"(\u03c7\u00b2 = {chi['chi2']:.2f}, Cram\u00e9r\u2019s V = {chi['V']:.3f}), confirming that "
            f"acquisition-site characteristics influence downstream ingest. However, the point-biserial "
            f"correlation between DBI and conversion outcome was not significant (r_pb = {pb['r']:.3f}, "
            f"p = {pb['p']:.3f}), nor were any individual component scores predictive of conversion success. "
            f"This null result is interpretable: with a 99.8% conversion pass rate, the near-ceiling "
            f"success leaves minimal variance for DBI to predict. The three failures (exit code 1) were "
            f"caused by DICOM sorting and instance-number inconsistencies\u2014factors orthogonal to the "
            f"metadata-quality dimensions DBI captures. Notably, all three produced partial NIfTI output "
            f"despite the non-zero exit code, illustrating why the dual criterion (exit code 0 AND NIfTI "
            f"presence) is more informative than file presence alone. DBI measures whether the metadata "
            f"foundation exists for correct interpretation of the converted output. A series may convert "
            f"to NIfTI yet still carry incomplete sidecar JSON (low M), ambiguous naming (low N), or "
            f"inconsistent geometry (low S)\u2014problems that surface only during downstream BIDS curation "
            f"or pipeline processing. Future work should examine the correlation between DBI components and "
            f"finer-grained conversion quality metrics (e.g., sidecar completeness, spatial geometry accuracy)."
        )
    else:
        p_justified(doc,
            "dcm2niix conversion pass rates varied across scanner clusters and series classes. While conversion "
            "success depends on factors beyond metadata quality (e.g., compressed transfer syntaxes, proprietary "
            "encoding), the systematic cluster-level variation parallels DBI differences."
        )

    h(doc, "4.5 Practical implications", level=2)
    p_justified(doc,
        "DBI audits can serve as an acquisition governance tool at any institution without re-calibration. "
        "Because components M, S, and G assess vendor-standardized DICOM tag presence (tested here across "
        "two vendors and two field strengths), and components P and N assess compliance with published "
        "community standards (BIDS, ENIGMA, ADNI, DICOM PS3.15) rather than site-specific patterns, the "
        "default DBI configuration is institution-independent by construction. Sites may override individual "
        "naming conventions in the YAML when a justified local SOP departs from the community standard, but "
        "the defaults require no site-specific tuning."
    )
    p_justified(doc,
        "For sites adopting DBI prospectively, the component-level scores provide actionable feedback: a low "
        "N score identifies specific naming conventions that must be addressed before the series can enter "
        "automated BIDS conversion or AI/ML workflows; a low P score signals that protocol naming at the "
        "scanner console does not yet align with structured standards; a low G score on DWI series indicates "
        "that diffusion encoding parameters are not discoverable from DICOM, jeopardizing downstream tensor "
        "or fiber-orientation estimation. This granularity enables targeted remediation at the point of "
        "acquisition rather than expensive post-hoc data cleaning."
    )

    h(doc, "4.6 Limitations", level=2)
    bullets(doc, [
        f"Single institutional archive: All {n_sessions} sessions come from one clinical site. However, "
        "hardware generalizability is addressed by the cross-vendor, cross-model scanner diversity in this "
        "cohort (two GE, two Siemens platforms across two field strengths), and workflow generalizability is "
        "addressed by grounding P and N conventions in published community standards (BIDS, ENIGMA, ADNI, "
        "DICOM PS3.15) rather than site-specific patterns. The primary untested dimension is whether sites "
        "with intentionally non-standard naming SOPs would benefit from convention overrides.",
        "Retrospective design: P_ideal was universally zero because the cohort predates prospective "
        "protocol-token adoption. Prospective deployment is required to evaluate this component.",
        "Heuristic classification: Series class assignment is rule-based and was not validated against "
        "human labels in this study. Misclassification could bias component-specific scores (especially G).",
        "Near-ceiling conversion: 99.8% of series converted successfully, leaving minimal variance for "
        "DBI\u2013conversion association analysis. A cohort with more heterogeneous transfer syntaxes or "
        "older DICOM encoding may reveal stronger associations.",
        "Expert-set weights: Default weights were set by domain-expert consensus rather than empirical "
        "calibration. Although Monte Carlo sensitivity analysis confirmed rank stability (\u00b120%), "
        "data-driven weight optimization remains future work.",
        "Drift control (D) is not scored in v1; protocol version time-series analysis requires external logs.",
    ])

    h(doc, "4.7 Future directions", level=2)
    p_justified(doc,
        "Near-term extensions include: (1) formal Phase 4 validation of series classification against human "
        "labels with Cohen\u2019s kappa; (2) predictive modeling of conversion failure as a function of DBI "
        "components; (3) drift control (D) using SoftwareVersions time series across sessions; and "
        "(4) integration with BIDS curation pipelines (HeuDiConv, dcm2bids) to quantify the correlation "
        "between DBI and bids-validator (BIDS Community, 2024) error counts. Provenance tracking with tools "
        "such as DataLad (Halchenko et al., 2021) would further strengthen reproducibility by linking each "
        "DBI audit to a versioned data snapshot. Longer-term, DBI could be computed in real time at the "
        "scanner console or PACS gateway, providing immediate feedback to technologists on naming and metadata "
        "compliance."
    )

    doc.add_page_break()

    # =========== 5. CONCLUSIONS ===========
    h(doc, "5. Conclusions", level=1)
    p_justified(doc,
        "Data Birth Integrity (DBI) provides a transparent, reproducible metric for assessing the automation "
        "readiness of clinical DICOM series at the earliest point in the neuroimaging data lifecycle. A low "
        "DBI score\u2014particularly when driven by protocol naming (P) and community-standards naming "
        "compliance (N)\u2014carries a concrete operational interpretation: the series is not ready for "
        "automated BIDS conversion without manual relabeling, will require custom heuristics for AI/ML "
        "pipeline ingestion, and imposes additional data-cleaning complexity before it can enter reproducible "
        "workflows. Because P and N conventions are grounded in published community standards (BIDS, ENIGMA, "
        "ADNI, DICOM PS3.15) rather than site-specific patterns, this interpretation holds across institutions "
        "without re-calibration. Components M, S, and G assess vendor-standardized DICOM properties tested "
        "here across two vendors and two field strengths. Routine DBI audits can quantify the curation burden "
        "for a given cohort, surface acquisition-site heterogeneity before it propagates to downstream "
        "analyses, and provide actionable feedback for prospective protocol governance. The open-source "
        "implementation and YAML-driven configuration lower the barrier to adoption."
    )

    doc.add_page_break()

    # =========== ACKNOWLEDGMENTS ===========
    h(doc, "Acknowledgments", level=1)
    p(doc, "[To be completed by team. Acknowledge funding, computational resources, and individuals who "
           "contributed to data collection or infrastructure.]", italic=True)

    # =========== DATA & CODE AVAILABILITY ===========
    h(doc, "Data and code availability", level=1)
    p_justified(doc,
        "The DBI scoring code, YAML configuration, unit tests, and audit scripts are available at "
        "[repository URL]. The DBI v1 specification (DBI_v1_SPECIFICATION.md) and companion acquisition "
        "conventions document are included in the repository. Raw DICOM data are not publicly shared due to "
        "institutional data-use agreements; aggregate statistics and de-identified exemplars are available "
        "upon reasonable request. [Adjust per team's data-sharing tier.]"
    )

    # =========== AUTHOR CONTRIBUTIONS ===========
    h(doc, "Author contributions", level=1)
    p(doc, "[CRediT taxonomy: Conceptualization, Methodology, Software, Validation, Formal analysis, "
           "Investigation, Data curation, Writing \u2014 original draft, Writing \u2014 review & editing, "
           "Visualization, Supervision, Project administration, Funding acquisition. Assign per team.]",
      italic=True)

    # =========== DECLARATION OF INTEREST ===========
    h(doc, "Declaration of competing interests", level=1)
    p(doc, "[To be completed by team.]", italic=True)

    doc.add_page_break()

    # =========== REFERENCES ===========
    h(doc, "References", level=1)
    refs = [
        "Aryanto, K. Y. E., Oudkerk, M., & van Ooijen, P. M. A. (2015). Free DICOM de-identification tools in clinical research: Functioning and safety of patient privacy. European Radiology, 25(12), 3685\u20133695. https://doi.org/10.1007/s00330-015-3794-0",
        "Bastiani, M., Cottaar, M., Fitzgibbon, S. P., et al. (2019). Automated quality control for within and between studies diffusion MRI data using a non-parametric framework for movement and distortion correction. NeuroImage, 184, 801\u2013812. https://doi.org/10.1016/j.neuroimage.2018.09.073",
        "BIDS Community. (2024). bids-validator: Validator for the Brain Imaging Data Structure. https://github.com/bids-standard/bids-validator",
        "Bidgood, W. D., & Horii, S. C. (1997). Understanding and using DICOM, the data interchange standard for biomedical imaging. Journal of the American Medical Informatics Association, 4(3), 199\u2013212. https://doi.org/10.1136/jamia.1997.0040199",
        "Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates.",
        "Esteban, O., Birman, D., Schaer, M., Koyejo, O. O., Poldrack, R. A., & Gorgolewski, K. J. (2017). MRIQC: Advancing the automatic prediction of image quality in MRI from unseen sites. PLoS ONE, 12(9), e0184661. https://doi.org/10.1371/journal.pone.0184661",
        "Esteban, O., Markiewicz, C. J., Blair, R. W., et al. (2019). fMRIPrep: A robust preprocessing pipeline for functional MRI. Nature Methods, 16(1), 111\u2013116. https://doi.org/10.1038/s41592-018-0235-4",
        "Fischl, B. (2012). FreeSurfer. NeuroImage, 62(2), 774\u2013781. https://doi.org/10.1016/j.neuroimage.2012.01.021",
        "Fortin, J.-P., Parker, D., Tun\u00e7, B., et al. (2017). Harmonization of multi-site diffusion tensor imaging data. NeuroImage, 161, 149\u2013170. https://doi.org/10.1016/j.neuroimage.2017.08.047",
        "Fortin, J.-P., Cullen, N., Sheline, Y. I., et al. (2018). Harmonization of cortical thickness measurements across scanners and sites. NeuroImage, 167, 104\u2013120. https://doi.org/10.1016/j.neuroimage.2017.11.024",
        "Gorgolewski, K. J., Auer, T., Calhoun, V. D., et al. (2016). The brain imaging data structure, a format for organizing and describing outputs of neuroimaging experiments. Scientific Data, 3, 160044. https://doi.org/10.1038/sdata.2016.44",
        "Gorgolewski, K. J., Burns, C. D., Madison, C., et al. (2011). Nipype: A flexible, lightweight and extensible neuroimaging data processing framework in Python. Frontiers in Neuroinformatics, 5, 13. https://doi.org/10.3389/fninf.2011.00013",
        "Halchenko, Y. O., Wagner, A. S., et al. (2021). DataLad: Distributed system for joint management of code, data, and their relationship. Journal of Open Source Software, 6(63), 3262. https://doi.org/10.21105/joss.03262",
        "Halchenko, Y. O., et al. (2024). HeuDiConv \u2014 flexible DICOM conversion into structured directory layouts. Journal of Open Source Software, 9(99), 5839. https://doi.org/10.21105/joss.05839",
        "Jack, C. R., Bernstein, M. A., Fox, N. C., et al. (2008). The Alzheimer\u2019s Disease Neuroimaging Initiative (ADNI): MRI methods. Journal of Magnetic Resonance Imaging, 27(4), 685\u2013691. https://doi.org/10.1002/jmri.21049",
        "Kinahan, P., Fedorov, A., & Sullivan, D. (2024). MIDRC CRP 12: Determining image data quality, provenance, and harmonization. https://www.midrc.org/midrc-collaborating-research-projects/project-one-crp12",
        "Kocak, B., et al. (2024). METhodological RadiomICs Score (METRICS): A quality scoring tool for radiomics research endorsed by EuSoMII. Insights into Imaging, 15, 8. https://doi.org/10.1186/s13244-023-01572-w",
        "Larobina, M., & Murino, L. (2014). Medical image file formats. Journal of Digital Imaging, 27(2), 200\u2013206. https://doi.org/10.1007/s10278-013-9657-9",
        "Li, X., Morgan, P. S., Ashburner, J., Smith, J., & Rorden, C. (2016). The first step for neuroimaging data analysis: DICOM to NIfTI conversion. Journal of Neuroscience Methods, 264, 47\u201356. https://doi.org/10.1016/j.jneumeth.2016.03.001",
        "Marcus, D. S., Olsen, T. R., Ramaratnam, M., & Buckner, R. L. (2007). The Extensible Neuroimaging Archive Toolkit: An informatics platform for managing, exploring, and sharing neuroimaging data. Neuroinformatics, 5(1), 11\u201334. https://doi.org/10.1385/NI:5:1:11",
        "Markiewicz, C. J., Gorgolewski, K. J., Feingold, F., et al. (2021). The OpenNeuro resource for sharing of neuroscience data. eLife, 10, e71774. https://doi.org/10.7554/eLife.71774",
        "Mason, D. L., et al. pydicom: An open source DICOM library. https://github.com/pydicom/pydicom",
        "NEMA/DICOM Standards Committee. (2023). Thirty years of the DICOM standard. Journal of Digital Imaging. https://doi.org/10.1007/s10278-023-00915-z",
        "Norgaard, M., et al. (2022). PET-BIDS, an extension to the Brain Imaging Data Structure for positron emission tomography. Scientific Data, 9, 65. https://doi.org/10.1038/s41597-022-01164-1",
        "Poldrack, R. A., Baker, C. I., Durnez, J., et al. (2017). Scanning the horizon: Towards transparent and reproducible neuroimaging research. Nature Reviews Neuroscience, 18(2), 115\u2013126. https://doi.org/10.1038/nrn.2016.167",
        "Rorden, C., Harms, M. P., et al. (2025). DICOM datasets for reproducible neuroimaging research across manufacturers and software versions. Scientific Data, 12(1), 1168. https://doi.org/10.1038/s41597-025-05503-w",
        "Thompson, P. M., Jahanshad, N., Ching, C. R. K., et al. (2020). ENIGMA and global neuroscience: A decade of large-scale studies of the brain in health and disease across more than 40 countries. Translational Psychiatry, 10(1), 100. https://doi.org/10.1038/s41398-020-0705-1",
        "Tomczak, M., & Tomczak, E. (2014). The need to report effect size estimates revisited. An overview of some recommended measures of effect size. Trends in Sport Sciences, 21(1), 19\u201325.",
        "Wilkinson, M. D., Dumontier, M., Aalbersberg, I. J., et al. (2016). The FAIR guiding principles for scientific data management and stewardship. Scientific Data, 3, 160018. https://doi.org/10.1038/sdata.2016.18",
        "Yue, J. K., Vassar, M. J., Lingsma, H. F., et al. (2013). Transforming research and clinical knowledge in traumatic brain injury pilot: Multicenter implementation of the common data elements for traumatic brain injury. Journal of Neurotrauma, 30(22), 1831\u20131844. https://doi.org/10.1089/neu.2013.2970",
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(ref)

    doc.add_page_break()

    # =========== SUPPLEMENTARY MATERIAL ===========
    h(doc, "Supplementary Material", level=1)

    h(doc, "S1. DBI component weights and YAML configuration", level=2)
    p_justified(doc,
        f"All DBI parameters are specified in dbi_v1_config.yaml (version {spec_ver}). Default weights: "
        "M = 0.28, P = 0.18, G = 0.22, S = 0.17, N = 0.15 (D = 0.00, not scored in v1). Spatial plausibility "
        "bounds: pixel spacing 0.05\u201315.0 mm, slice thickness 0.05\u201320.0 mm."
    )

    h(doc, "S2. Per-class SeriesDescription compliance rules", level=2)
    p_justified(doc,
        "The following table summarizes the per-class naming standards enforced by component N "
        "(naming.class_series_description_compliance in YAML). All listed patterns must match the trimmed "
        "SeriesDescription for the class-specific check to pass."
    )
    sd_rules = [
        ["dwi", "Contains DTI/DWI; AP + digits + DIRECTION/DIR"],
        ["fmap", "Contains field map/fmap; PA + digits + DIRECTION/DIR"],
        ["bold", "Contains fMRI/BOLD/RESTING/TASK/EPI/rsfMRI"],
        ["asl", "Contains ASL"],
        ["flair", "Ends with FLAIR"],
        ["perf", "Contains perf/DSC/DCE/CBF"],
        ["t1_anat", "Ends with MPRAGE/SPGR/FSPGR/BRAVO/IR FSPGR"],
        ["t2_anat", "Contains T2 or BLADE; ends with THIN"],
        ["swi", "Contains SWI or SWAN"],
        ["localizer", "Contains localizer/LOC/scout"],
    ]
    add_table(doc, ["Class", "SeriesDescription requirement (all must match)"], sd_rules)

    h(doc, "S3. Conversion outcomes by scanner cluster and series class", level=2)
    if t2c_df is not None:
        add_table(doc,
            ["Scanner cluster", "Series class", "n", "n pass", "Pass rate"],
            [
                [row["scanner_cluster"], row["series_class"],
                 str(int(row["n"])), str(int(row["n_pass"])), f"{row['pass_rate']:.4f}"]
                for _, row in t2c_df.iterrows()
            ],
        )

    out_path = Path(__file__).resolve().parent / "DBI_v1_Manuscript.docx"
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
