#!/usr/bin/env python3
"""
Generate MR acquisition SOP-style DOCX: automation, AI/cleaning, and BIDS-inspired naming conventions.

Reads cidur_dbi/dbi_v1_config.yaml: naming.automation_conventions, class_series_description_compliance,
derived_scan_naming. Output: MR_Acquisition_Automation_and_BIDS_Inspired_Naming_Conventions.docx
"""

from __future__ import annotations

from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
CONFIG = ROOT / "cidur_dbi" / "dbi_v1_config.yaml"
OUT = ROOT / "MR_Acquisition_Automation_and_BIDS_Inspired_Naming_Conventions.docx"

# Plain-language acquisition guidance for §6.1 (keys = YAML class names under class_series_description_compliance).
COMPLIANCE_ACQ_GUIDANCE: dict[str, str] = {
    "dwi": (
        "Label as diffusion (DTI/DWI/dwi/dti). Include **AP** phase-encoding with a digit and explicit "
        "direction wording (e.g. DIRECTIONS, DIR, direction)."
    ),
    "fmap": (
        "Label as field map / fmap / GRE field map (including FMRI…field style). Include **PA** encoding "
        "with a digit and explicit direction wording."
    ),
    "bold": (
        "Include functional MRI intent: fMRI, fmri, BOLD, resting, task, EPI, or rsfMRI in SeriesDescription."
    ),
    "asl": (
        "Include **ASL** in SeriesDescription."
    ),
    "flair": (
        "End SeriesDescription with **FLAIR** or **flair** (trailing spaces allowed)."
    ),
    "perf": (
        "Include perfusion-related wording: perf, DSC, DCE, or CBF."
    ),
    "t1_anat": (
        "End SeriesDescription with a structural T1 sequence token: MPRAGE, SPGR, FSPGR, BRAVO, or IR FSPGR."
    ),
    "t2_anat": (
        "Include **T2** or **BLADE**; end with **THIN** (thin-section T2 convention)."
    ),
    "swi": (
        "Include **SWI** or **SWAN** in SeriesDescription."
    ),
    "localizer": (
        "Use localizer, Localizer, LOC, scout, or Scout naming per site SOP."
    ),
}


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    conventions: list[dict] = []
    scan_pat = "^[0-9]+-"
    proto_pat = "PROTO_[A-Z0-9_]+"
    cfg: dict = {}
    naming: dict = {}
    if CONFIG.is_file():
        with open(CONFIG, encoding="utf-8") as f:
            cfg = yaml.safe_load(f) or {}
        naming = cfg.get("naming") or {}
        conventions = naming.get("automation_conventions") or []
        scan_pat = naming.get("scan_folder_pattern", scan_pat)
        proto = cfg.get("protocol_token") or {}
        proto_pat = proto.get("pattern", proto_pat)
    ver = cfg.get("version", "1.0.1")
    class_sd_compliance = naming.get("class_series_description_compliance") or {}
    derived_naming = naming.get("derived_scan_naming") or {}

    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "MR imaging acquisition: automation-ready and BIDS-inspired naming conventions"
    )
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph()
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.add_run(
        "For data cleaning, AI pipelines, secondary analysis, and DBI v1 naming compliance (component N)"
    ).italic = True
    doc.add_paragraph()
    p(doc, f"Companion config version: {ver} (`cidur_dbi/dbi_v1_config.yaml`). Not a substitute for IRB or vendor safety review.")

    h(doc, "1. Purpose and audience", level=1)
    p(
        doc,
        "This document translates machine-checkable naming rules used in automated DICOM audits into "
        "practical guidance for MR acquisition and imaging operations. It supports:",
    )
    bullets(
        doc,
        [
            "Consistent SeriesDescription and ProtocolName text on the scanner console.",
            "Archive or PACS export conventions (e.g. numbered series folders) where applicable.",
            "Downstream automation: conversion (e.g. dcm2niix), scripting, machine learning, and BIDS-oriented curation.",
        ],
    )

    h(doc, "2. How this ties to DBI and audit software", level=1)
    p(
        doc,
        "Data Birth Integrity (DBI) v1 includes a naming subscore **N** (naming compliance). Part of **N** is "
        "evaluated on a combined string built from the scan folder basename (when present), DICOM "
        "**Series Description (0008,103E)**, and **Protocol Name (0018,1030)**. Optional checks listed in "
        "`naming.automation_conventions` reward patterns aligned with automation and BIDS-adjacent practice. "
        "If both description fields are empty and no folder name is available, only legacy structural checks apply.",
    )
    bullets(
        doc,
        [
            "Full technical definition: `cidur_dbi/DBI_v1_SPECIFICATION.md` §4.5.",
            "Live parameters and regexes: `cidur_dbi/dbi_v1_config.yaml` → `naming`.",
        ],
    )

    h(doc, "3. “BIDS-inspired” vs full BIDS", level=1)
    p(
        doc,
        "The **Brain Imaging Data Structure (BIDS)** specifies filenames, folder hierarchy, and JSON sidecars for "
        "research datasets. Clinical DICOM exports are **not** BIDS layouts. Here, “BIDS-inspired” means:",
    )
    bullets(
        doc,
        [
            "Using **entity-like tokens** in text (e.g. run-, acq-, dir-, echo-) where they help humans and scripts parse intent.",
            "Using **modality shorthand** familiar in neuroimaging (e.g. T1w, FLAIR, dwi, bold) when appropriate in descriptions.",
            "Avoiding claims of **full BIDS compliance** from the scanner console alone—BIDS is validated on the **curated dataset**, not raw DICOM naming.",
        ],
    )

    h(doc, "4. Archive and session naming (legacy exports)", level=1)
    p(
        doc,
        f"Many institutional exports use **indexed scan folder names** matching a pattern such as `{scan_pat}` "
        "(digits, hyphen, then description). This supports ordering and scripting. Acquisition teams should "
        "coordinate with PACS/XNAT administrators so exported folder basenames remain stable and parseable.",
    )

    h(doc, "5. Prospective protocol tokens (P_ideal in DBI)", level=1)
    p(
        doc,
        f"DBI component **P** includes an “ideal” subscore when SeriesDescription or ProtocolName matches a "
        f"site-defined regex (example in config: `{proto_pat}`). Replace the example with your real institutional "
        "token when rolling out prospective studies. Legacy data may score low on this axis by design.",
    )

    h(doc, "6. Automation and cleaning-oriented conventions (configurable)", level=1)
    p(
        doc,
        "The following conventions are implemented as optional **binary** checks in DBI **N** (each passes if the "
        "regex finds a match in the combined naming text). Sites may disable or edit entries in YAML. "
        "Acquisition guidance summarizes intent; the YAML regex is authoritative for scoring.",
    )

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "ID"
    hdr[1].text = "Goal (automation / AI / cleaning)"
    hdr[2].text = "Acquisition guidance"
    hdr[3].text = "Implementation note"

    guidance_map = {
        "no_control_chars": (
            "Do not paste multi-line text or hidden characters into SeriesDescription/ProtocolName; "
            "use single-line, printable text suitable for logs and CSV.",
        ),
        "bids_style_entity_tokens": (
            "Where protocol design allows, include parseable tokens such as run-1, acq-highres, dir-AP, echo-1 "
            "in descriptions so curators can map to BIDS entities later.",
        ),
        "modality_or_contrast_lexicon": (
            "Include a clear modality or contrast token (e.g. MPRAGE, FLAIR, DWI, BOLD, ASL, field map) in "
            "SeriesDescription or ProtocolName for automated series role detection.",
        ),
        "token_separators": (
            "Prefer hyphens or underscores between tokens (e.g. 3D_T2_FLAIR, T1w_MPRAGE) rather than long "
            "unbroken strings or ambiguous spaces.",
        ),
    }

    for conv in conventions:
        cid = conv.get("id", "")
        rat = conv.get("rationale", "")
        en = conv.get("enabled", True)
        rx = conv.get("regex", "")
        acq_g = guidance_map.get(
            cid,
            ("Align text with institutional SOP and DBI YAML rationale for this id.",),
        )
        if isinstance(acq_g, tuple):
            acq_g = acq_g[0]
        impl = f"Regex (Python): {rx}\nEnabled in config: {en}\nSpec rationale: {rat}"
        row = tbl.add_row().cells
        row[0].text = cid
        row[1].text = rat
        row[2].text = acq_g
        row[3].text = impl

    if not conventions:
        row = tbl.add_row().cells
        row[0].text = "—"
        row[1].text = "No rows loaded from YAML"
        row[2].text = str(CONFIG)
        row[3].text = "Check path and naming.automation_conventions"

    h(doc, "6.1 Per-class SeriesDescription standards (DBI N)", level=2)
    p(
        doc,
        "These rows are the **naming standards** acquisition should use for DICOM **Series Description (0008,103E)**. "
        "The audit first assigns a **series class** (broader legacy patterns in `classification_rules`); then, for "
        "that class, **every** regex below must match **trimmed SeriesDescription alone** or the class-specific part "
        "of naming compliance fails. Empty SeriesDescription fails. Authoritative source: "
        "`naming.class_series_description_compliance` in `dbi_v1_config.yaml`.",
    )

    std_tbl = doc.add_table(rows=1, cols=3)
    std_tbl.style = "Table Grid"
    sh = std_tbl.rows[0].cells
    sh[0].text = "Series class"
    sh[1].text = "Acquisition standard (plain language)"
    sh[2].text = "Machine checks — all must match (regex on SeriesDescription)"

    if not class_sd_compliance:
        row = std_tbl.add_row().cells
        row[0].text = "—"
        row[1].text = "No class_series_description_compliance loaded"
        row[2].text = str(CONFIG)
    else:
        for class_name, spec in class_sd_compliance.items():
            patterns = spec.get("all_match") or []
            acq = COMPLIANCE_ACQ_GUIDANCE.get(
                class_name,
                "Satisfy every machine check in this row on SeriesDescription (see YAML for rationale).",
            )
            tech = "\n".join(patterns) if patterns else "(no all_match patterns in YAML)"
            row = std_tbl.add_row().cells
            row[0].text = str(class_name)
            row[1].text = acq
            row[2].text = tech

    h(doc, "6.2 Derived and reformatted series", level=2)
    p(
        doc,
        "When **folder basename + SeriesDescription + ProtocolName** contains a **derivative marker** as a "
        "delimiter-bounded token (same idea as `naming.derivative_tokens`), **SeriesDescription** should end with "
        "**reformat** or **derived** so pipelines can tell curated outputs from raw acquisitions. "
        "Configured under `naming.derived_scan_naming` in YAML.",
    )
    markers = derived_naming.get("markers") or []
    suff = str(derived_naming.get("series_description_suffix_regex", ""))
    den = derived_naming.get("enabled", True)
    der_tbl = doc.add_table(rows=1, cols=2)
    der_tbl.style = "Table Grid"
    dh = der_tbl.rows[0].cells
    dh[0].text = "Setting"
    dh[1].text = "Value"
    for label, val in (
        ("Enabled in config", "true" if den else "false"),
        ("Derivative markers (combined text)", ", ".join(str(m) for m in markers) if markers else "—"),
        ("SeriesDescription must match (suffix)", suff if suff else "—"),
    ):
        dr = der_tbl.add_row().cells
        dr[0].text = label
        dr[1].text = val

    h(doc, "7. DWI and derived maps", level=1)
    p(
        doc,
        "For diffusion, clearly distinguish **raw encoding series** from **derived** maps (ADC, FA, TRACE, etc.) "
        "in names and descriptions so pipelines do not treat scalars as raw DWI. DBI uses configurable "
        "derivative_tokens for gradient scoring; acquisition should avoid ambiguous labels that mix raw and derived intent.",
    )

    h(doc, "8. Governance", level=1)
    bullets(
        doc,
        [
            "Update `dbi_v1_config.yaml` when conventions change; cite config version or hash in study protocols and methods.",
            "Re-run `run_audit.py` after YAML changes to refresh cohort statistics.",
            "This document is descriptive SOP text; enforceability is institutional policy.",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
