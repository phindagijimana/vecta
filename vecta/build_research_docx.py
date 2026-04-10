#!/usr/bin/env python3
"""Generate Research Project Documentation as .docx."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = Document()

    # Title page style title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Research Project Documentation")
    r.bold = True
    r.font.size = Pt(18)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run(
        "AI-Derived Acquisition Standards for Diffusion MRI (DWI) "
        "and Data Birth Integrity (DBI)"
    )
    sr.bold = True
    sr.font.size = Pt(14)

    doc.add_paragraph()

    # 1 Project Title
    add_heading(doc, "1. Project Title", level=1)
    doc.add_paragraph(
        "AI-Derived Acquisition Contracts for Diffusion MRI: Improving Biomarker "
        "Stability and Model Generalization through Data Birth Integrity"
    )

    # 2 Executive Summary
    add_heading(doc, "2. Executive Summary", level=1)
    doc.add_paragraph(
        "Diffusion-weighted imaging (DWI) is widely used in neuroscience and clinical "
        "research, yet suffers from high acquisition variability, leading to:"
    )
    add_bullets(
        doc,
        [
            "Poor reproducibility",
            "Limited biomarker reliability",
            "Reduced cross-site generalization",
            "Scarcity of reusable public datasets",
        ],
    )
    doc.add_paragraph("This project proposes a new framework:")
    add_bullets(
        doc,
        [
            "Deriving acquisition-level standards from the structural requirements of "
            "AI and automation systems.",
        ],
    )
    doc.add_paragraph("We introduce:")
    add_bullets(
        doc,
        [
            "AI-Ready Acquisition Schema (AI-RAS) for DWI",
            "Data Birth Integrity (DBI) as a measurable structural quality metric",
            "A validation framework to enforce acquisition constraints",
            "An empirical evaluation of impact on biomarkers and AI models",
        ],
    )
    doc.add_paragraph(
        "The project will be implemented at an imaging center and evaluated through "
        "controlled experiments."
    )

    # 3 Problem Statement
    add_heading(doc, "3. Problem Statement", level=1)
    doc.add_paragraph("Current DWI datasets exhibit:")
    add_bullets(
        doc,
        [
            "Inconsistent b-values",
            "Variable gradient directions",
            "Missing or incomplete metadata",
            "Scanner-dependent variability",
            "Protocol drift across time and sites",
        ],
    )
    doc.add_paragraph("These inconsistencies:")
    add_bullets(
        doc,
        [
            "Increase harmonization burden",
            "Reduce reproducibility",
            "Introduce bias into AI models",
            "Limit dataset reuse",
        ],
    )
    doc.add_paragraph("Existing standards such as:")
    add_bullets(doc, ["DICOM", "Brain Imaging Data Structure"])
    doc.add_paragraph("do not enforce acquisition-level consistency.")

    # 4 Hypothesis
    add_heading(doc, "4. Research Hypothesis", level=1)
    doc.add_paragraph(
        "Enforcing AI-derived acquisition constraints at data creation improves diffusion "
        "biomarker reproducibility and machine learning generalization compared to "
        "standard post-hoc harmonization approaches."
    )

    # 5 Conceptual Framework
    add_heading(doc, "5. Conceptual Framework", level=1)

    add_heading(doc, "5.1 Data Generation Pipeline", level=2)
    doc.add_paragraph("Reality → Measurement → Encoding → Dataset → AI Model")
    doc.add_paragraph(
        "Errors introduced at the measurement and encoding stages propagate downstream."
    )

    add_heading(doc, "5.2 Key Definitions", level=2)
    defs = [
        (
            "Data Birth Integrity (DBI)",
            "A quantitative measure of structural correctness at data creation.",
        ),
        (
            "AI-Ready Acquisition Schema (AI-RAS)",
            "A set of acquisition constraints derived from AI pipeline requirements.",
        ),
        (
            "Acquisition Contract",
            "A versioned specification defining required acquisition conditions.",
        ),
        (
            "Protocol Drift Index (PDI)",
            "A measure of deviation from defined acquisition protocol across time or sites.",
        ),
    ]
    for term, desc in defs:
        p = doc.add_paragraph()
        p.add_run(term + ". ").bold = True
        p.add_run(desc)

    # 6 AI-Derived DWI Acquisition Standards
    add_heading(doc, "6. AI-Derived DWI Acquisition Standards", level=1)
    doc.add_paragraph(
        "These standards are derived from requirements of diffusion modeling, "
        "tractography, and machine learning pipelines."
    )

    add_heading(doc, "6.1 Multi-Shell Requirement", level=2)
    doc.add_paragraph("Must include:")
    add_bullets(
        doc,
        [
            "b = 0",
            "≥1 diffusion shell ≥1000",
            "Prefer multi-shell acquisition (e.g., 0, 1000, 2000)",
        ],
    )

    add_heading(doc, "6.2 Gradient Direction Requirement", level=2)
    doc.add_paragraph("Minimum:")
    add_bullets(
        doc,
        [
            "≥30 directions (single-shell)",
            "≥60 directions (multi-shell)",
        ],
    )

    add_heading(doc, "6.3 Gradient Table Integrity", level=2)
    doc.add_paragraph("Mandatory presence of:")
    add_bullets(doc, ["bvals", "bvecs"])
    doc.add_paragraph("Alignment and consistency checks required.")

    add_heading(doc, "6.4 Spatial Resolution Constraint", level=2)
    doc.add_paragraph("Target voxel size: ~2 mm isotropic (acceptable range defined).")

    add_heading(doc, "6.5 Phase Encoding Metadata", level=2)
    doc.add_paragraph("Required fields:")
    add_bullets(
        doc,
        [
            "PhaseEncodingDirection",
            "TotalReadoutTime",
            "Fieldmap availability (if applicable)",
        ],
    )

    add_heading(doc, "6.6 Protocol Versioning", level=2)
    doc.add_paragraph('Each acquisition must include: protocol_version = "DWI_V1.0"')

    add_heading(doc, "6.7 Naming Consistency", level=2)
    doc.add_paragraph("Example:")
    add_bullets(
        doc,
        [
            "sub-001_ses-01_dwi.nii.gz",
            "sub-001_ses-01_dwi.bval",
            "sub-001_ses-01_dwi.bvec",
        ],
    )

    add_heading(doc, "6.8 Mandatory Metadata Fields", level=2)
    add_bullets(
        doc,
        [
            "b-values",
            "gradient directions",
            "voxel size",
            "TE/TR",
            "scanner vendor",
            "field strength",
        ],
    )

    # 7 DBI Metric
    add_heading(doc, "7. Data Birth Integrity (DBI) Metric", level=1)

    add_heading(doc, "7.1 DBI Components", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Description"
    rows = [
        ("Metadata completeness", "Required fields present"),
        ("Protocol conformity", "Matches acquisition contract"),
        ("Gradient integrity", "Valid diffusion encoding"),
        ("Spatial consistency", "Voxel size within range"),
        ("Naming compliance", "Standardized naming"),
        ("Drift control", "Protocol adherence over time"),
    ]
    for comp, desc in rows:
        row = table.add_row().cells
        row[0].text = comp
        row[1].text = desc

    add_heading(doc, "7.2 Example DBI Formula", level=2)
    doc.add_paragraph("DBI = w1*M + w2*P + w3*G + w4*S + w5*N + w6*D")
    doc.add_paragraph("Where:")
    add_bullets(
        doc,
        [
            "M = metadata completeness",
            "P = protocol conformity",
            "G = gradient integrity",
            "S = spatial consistency",
            "N = naming compliance",
            "D = drift adherence",
        ],
    )

    # 8 Methodology
    add_heading(doc, "8. Methodology", level=1)

    add_heading(doc, "8.1 Study Design", level=2)
    doc.add_paragraph("Two conditions:")
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    h2 = t2.rows[0].cells
    h2[0].text = "Condition A"
    h2[1].text = "Condition B"
    r2 = t2.add_row().cells
    r2[0].text = "Standard acquisition"
    r2[1].text = "AI-RAS enforced acquisition"

    add_heading(doc, "8.2 Data Collection", level=2)
    add_bullets(
        doc,
        [
            "Retrospective dataset (baseline)",
            "Prospective dataset (after enforcement)",
        ],
    )

    add_heading(doc, "8.3 Metrics", level=2)
    for sub, items in [
        (
            "Structural Metrics",
            [
                "% missing metadata",
                "Naming inconsistency rate",
                "DBI score",
            ],
        ),
        (
            "Operational Metrics",
            [
                "Cleaning time (hours)",
                "Pipeline failure rate",
                "Reprocessing frequency",
            ],
        ),
        (
            "AI Metrics",
            [
                "Model AUC",
                "Cross-site performance variance",
                "Calibration stability",
            ],
        ),
        (
            "Biomarker Metrics",
            [
                "FA reproducibility",
                "MD stability",
                "Tractography consistency",
            ],
        ),
    ]:
        p = doc.add_paragraph()
        p.add_run(sub + ":").bold = True
        for it in items:
            doc.add_paragraph(it, style="List Bullet")

    add_heading(doc, "8.4 Statistical Analysis", level=2)
    add_bullets(
        doc,
        [
            "t-tests for group comparison",
            "Variance analysis",
            "Effect size calculation",
            "Correlation between DBI and model performance",
        ],
    )

    # 9 Implementation Plan
    add_heading(doc, "9. Implementation Plan", level=1)
    phases = [
        (
            "Phase 1 — Baseline Analysis",
            [
                "Analyze historical DWI datasets",
                "Compute DBI and PDI",
            ],
        ),
        (
            "Phase 2 — Schema Definition",
            [
                "Define AI-RAS for DWI",
                "Implement validation rules",
            ],
        ),
        (
            "Phase 3 — Pilot Deployment",
            [
                "Apply at imaging center",
                "Introduce validation at export stage",
            ],
        ),
        (
            "Phase 4 — Evaluation",
            [
                "Compare before vs after",
                "Analyze metrics",
            ],
        ),
    ]
    for title, bullets in phases:
        add_heading(doc, title, level=2)
        add_bullets(doc, bullets)

    # 10 Expected Outcomes
    add_heading(doc, "10. Expected Outcomes", level=1)
    for title, bullets in [
        (
            "Structural Improvements",
            ["Increased metadata completeness", "Reduced protocol drift"],
        ),
        (
            "Operational Improvements",
            ["Reduced cleaning time", "Fewer pipeline failures"],
        ),
        (
            "AI Improvements",
            ["Reduced model variance", "Improved generalization"],
        ),
        (
            "Scientific Impact",
            ["More reproducible biomarkers", "Increased dataset reusability"],
        ),
    ]:
        p = doc.add_paragraph()
        p.add_run(title + ":").bold = True
        add_bullets(doc, bullets)

    # 11 Contributions
    add_heading(doc, "11. Potential Contributions", level=1)
    for title, bullets in [
        (
            "Scientific Contributions",
            [
                "Formal definition of DBI",
                "Empirical evidence linking acquisition structure to AI performance",
                "Framework for acquisition-level standardization",
            ],
        ),
        (
            "Technical Contributions",
            [
                "AI-derived acquisition schema",
                "Validation framework",
                "Drift monitoring approach",
            ],
        ),
        (
            "Practical Contributions",
            [
                "Improved imaging workflows",
                "Reduced harmonization burden",
                "Better dataset sharing",
            ],
        ),
    ]:
        p = doc.add_paragraph()
        p.add_run(title + ":").bold = True
        add_bullets(doc, bullets)

    # 12 Limitations
    add_heading(doc, "12. Limitations", level=1)
    add_bullets(
        doc,
        [
            "Limited control over scanner hardware",
            "Site-specific variability",
            "Adoption resistance",
            "Need for multi-site validation",
        ],
    )

    # 13 Future Work
    add_heading(doc, "13. Future Work", level=1)
    add_bullets(
        doc,
        [
            "Extend to other imaging modalities",
            "Multi-site validation studies",
            "Integration with clinical data models (e.g., OMOP Common Data Model)",
            "Expansion into certification framework",
        ],
    )

    # 14 Broader Impact
    add_heading(doc, "14. Broader Impact", level=1)
    doc.add_paragraph("This work contributes to:")
    add_bullets(
        doc,
        [
            "Reproducible science",
            "Data-centric AI",
            "FAIR data principles",
            "Global data accessibility",
        ],
    )

    # Final Statement
    add_heading(doc, "Final Statement", level=1)
    doc.add_paragraph(
        "This project reframes standardization: from post-hoc harmonization to "
        "preemptive structural enforcement at data creation."
    )
    doc.add_paragraph(
        "By grounding acquisition standards in AI requirements, this work aims to:"
    )
    add_bullets(
        doc,
        [
            "Improve reliability of scientific data",
            "Enable scalable AI systems",
            "Reduce structural inefficiencies in data pipelines",
        ],
    )

    out = "/mnt/nfs/home/urmc-sh.rochester.edu/pndagiji/Documents/vecta/Research_Project_Documentation_DWI_DBI.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
